"""Generate the International Dual Momentum one-pager Word doc."""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import os

SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
MAROON = RGBColor(0x80, 0x00, 0x00)
BLACK = RGBColor(0, 0, 0)
GRAY = RGBColor(0x44, 0x44, 0x44)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
LIGHT_GRAY = RGBColor(0xF2, 0xF2, 0xF2)


def set_cell_shading(cell, color_hex):
    shading = cell._element.get_or_add_tcPr()
    s = shading.makeelement(qn("w:shd"), {
        qn("w:val"): "clear",
        qn("w:color"): "auto",
        qn("w:fill"): color_hex,
    })
    shading.append(s)


def add_styled_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"

    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        run.bold = True
        run.font.size = Pt(8)
        run.font.color.rgb = WHITE
        set_cell_shading(cell, "800000")

    for r_idx, row_data in enumerate(rows):
        for c_idx, val in enumerate(row_data):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = ""
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(str(val))
            run.font.size = Pt(8)
            run.font.color.rgb = BLACK
            if r_idx % 2 == 1:
                set_cell_shading(cell, "F2F2F2")

    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Inches(w)

    return table


def build():
    doc = Document()

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(9)
    style.paragraph_format.space_after = Pt(2)
    style.paragraph_format.space_before = Pt(0)

    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.4)
        section.left_margin = Inches(0.6)
        section.right_margin = Inches(0.6)

    # Header
    p = doc.add_paragraph()
    run = p.add_run("POTOMAC FUND MANAGEMENT")
    run.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = MAROON
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run2 = p.add_run("\t\t\t\t\t\t         March 2026")
    run2.font.size = Pt(10)
    run2.font.color.rgb = GRAY

    p = doc.add_paragraph()
    run = p.add_run("International Dual Momentum: Single-Country ETF Rotation")
    run.bold = True
    run.font.size = Pt(12)
    run.font.color.rgb = BLACK

    p = doc.add_paragraph()
    run = p.add_run("Research Concept — Preliminary Results, Requires Further Analysis")
    run.italic = True
    run.font.size = Pt(9)
    run.font.color.rgb = GRAY

    # Divider
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run("_" * 95)
    run.font.color.rgb = MAROON
    run.font.size = Pt(6)

    # Concept
    p = doc.add_paragraph()
    run = p.add_run("THE CONCEPT")
    run.bold = True
    run.font.size = Pt(10)
    run.font.color.rgb = MAROON

    doc.add_paragraph(
        "Apply Gary Antonacci\u2019s dual momentum framework to a universe of 16 liquid, "
        "single-country international ETFs (all iShares MSCI). Each month, rank countries by "
        "trailing 12-month total return (relative momentum), select the top 7, equal-weight at "
        "~14.3% each, and apply an absolute momentum filter: any country whose return fails the "
        "cash hurdle flips to SGOV. The strategy is 100% international \u2014 no US equity exposure."
    )

    # Universe
    p = doc.add_paragraph()
    run = p.add_run("UNIVERSE (16 ETFs)")
    run.bold = True
    run.font.size = Pt(10)
    run.font.color.rgb = MAROON

    p = doc.add_paragraph()
    run = p.add_run("Developed (9): ")
    run.bold = True
    run.font.size = Pt(8.5)
    run2 = p.add_run("EWJ Japan \u2022 EWG Germany \u2022 EWU UK \u2022 EWC Canada \u2022 "
                      "EWA Australia \u2022 EWQ France \u2022 EWL Switzerland \u2022 EWP Spain \u2022 EWI Italy")
    run2.font.size = Pt(8.5)

    p = doc.add_paragraph()
    run = p.add_run("Emerging (7):  ")
    run.bold = True
    run.font.size = Pt(8.5)
    run2 = p.add_run("EWT Taiwan \u2022 EWZ Brazil \u2022 INDA India \u2022 FXI China \u2022 "
                      "EWY S. Korea \u2022 EWW Mexico \u2022 EWH Hong Kong")
    run2.font.size = Pt(8.5)

    # Preliminary backtest
    p = doc.add_paragraph()
    run = p.add_run("PRELIMINARY BACKTEST (Jan 2004 \u2013 Feb 2026)")
    run.bold = True
    run.font.size = Pt(10)
    run.font.color.rgb = MAROON

    p = doc.add_paragraph()
    run = p.add_run("Five absolute momentum (go-to-cash) triggers were tested. Top-line results:")
    run.font.size = Pt(9)

    headers = ["Trigger", "CAGR", "Max DD", "Sharpe", "Sortino", "Avg Countries", "% Invested"]
    rows = [
        ["A: Classic (12m > BIL)", "2.5%", "-33.8%", "0.02", "0.02", "4.2", "61%"],
        ["B: Dual (12m>0 & 6m>0)", "6.9%", "-30.1%", "0.36", "0.51", "4.9", "70%"],
        ["C: Composite (avg 1/3/6/12m)", "9.2%", "-27.3%", "0.49", "0.75", "5.8", "82%"],
        ["D: Aggregate (EFA > BIL)", "3.0%", "-28.9%", "0.05", "0.05", "3.5", "49%"],
        ["E: Breadth (>50% > BIL)", "3.6%", "-28.9%", "0.10", "0.10", "3.4", "48%"],
    ]
    add_styled_table(doc, headers, rows, [2.2, 0.6, 0.7, 0.65, 0.65, 1.0, 0.8])

    doc.add_paragraph("")

    headers2 = ["Benchmark", "CAGR", "Max DD", "Sharpe", "Volatility"]
    rows2 = [
        ["EFA (EAFE buy-hold)", "6.8%", "-57.4%", "0.30", "16.6%"],
        ["ACWX (All-World ex-US)", "4.4%", "-56.2%", "0.16", "17.9%"],
        ["SPY (S&P 500 buy-hold)", "10.5%", "-50.8%", "0.56", "14.5%"],
    ]
    add_styled_table(doc, headers2, rows2, [2.2, 0.6, 0.7, 0.65, 0.8])

    # Current holdings
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    run = p.add_run("CURRENT SIGNAL (as of Jan 2026)")
    run.bold = True
    run.font.size = Pt(10)
    run.font.color.rgb = MAROON

    doc.add_paragraph(
        "Top 7: EWY (S. Korea), EWP (Spain), EWW (Mexico), EWZ (Brazil), EWI (Italy), "
        "EWH (Hong Kong), EWT (Taiwan). All 7 passing absolute momentum \u2014 0% cash. "
        "Themes: European value rotation + EM momentum. Heavy non-US developed Europe tilt."
    )

    # What needs more work
    p = doc.add_paragraph()
    run = p.add_run("AREAS REQUIRING FURTHER ANALYSIS")
    run.bold = True
    run.font.size = Pt(10)
    run.font.color.rgb = MAROON

    items = [
        ("Trigger C looks best but may be overfit.", " The composite trigger (avg of 1/3/6/12m) "
         "has 4 implicit parameters vs. Antonacci\u2019s single 12-month lookback. Need out-of-sample "
         "validation and sub-period stability testing before adopting."),
        ("BIL start date creates a problem.", " BIL only starts May 2007. Triggers A/D/E use BIL "
         "as the hurdle and show 0% returns 2004\u20132006 because the absolute momentum filter "
         "can\u2019t fire without cash benchmark data. Need to splice in 3-month T-bill rates from FRED "
         "pre-2007 for a clean backtest."),
        ("Transaction costs and slippage not modeled.", " Monthly rebalancing of 7 country ETFs "
         "generates turnover. Tier 3 ETFs (EWQ, EWP, EWI) have wider spreads. Need to model "
         "realistic execution costs, especially at scale."),
        ("Currency effects embedded but unexamined.", " All returns are USD-denominated. A strong "
         "dollar regime (like 2014\u20132015) crushes international returns regardless of local "
         "momentum. Should we add a USD filter or currency-hedged variants?"),
        ("No tail-risk protection.", " Max drawdown of -27% on the best trigger is better than "
         "EFA\u2019s -57%, but still significant. Consider adding a CAOS-style convexity overlay "
         "or VIX regime filter, consistent with the existing Potomac defensive architecture."),
        ("Concentration risk in current signal.", " 5 of 7 current holdings are in just two "
         "themes (European value + EM). Correlation within the basket during a risk-off event "
         "could be high. Need to stress-test the portfolio under 2008/2020 shock scenarios."),
        ("Integration with existing Potomac strategies.", " How does this complement or overlap "
         "with the current Bull Bear signal and CRTOX/CRTPX allocations? Need correlation "
         "analysis against the existing book."),
    ]

    for bold_part, rest in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(1)
        run = p.add_run(bold_part)
        run.bold = True
        run.font.size = Pt(8.5)
        run2 = p.add_run(rest)
        run2.font.size = Pt(8.5)

    # Bottom line
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    run = p.add_run("_" * 95)
    run.font.color.rgb = MAROON
    run.font.size = Pt(6)

    p = doc.add_paragraph()
    run = p.add_run("Bottom Line: ")
    run.bold = True
    run.font.size = Pt(9)
    run.font.color.rgb = MAROON
    run2 = p.add_run(
        "The concept has merit \u2014 country momentum is a well-documented factor (Asness, Moskowitz, "
        "Pedersen 2013), and the absolute momentum filter materially reduces drawdowns vs. passive "
        "international. But the backtest has data gaps, the best trigger may be overfit, and "
        "implementation details (costs, currency, tail risk, portfolio integration) are unresolved. "
        "This is a research starting point, not a tradeable strategy yet. Recommend dedicating "
        "next session to fixing the BIL data issue and running sub-period robustness checks."
    )
    run2.font.size = Pt(9)

    path = os.path.join(SCRIPT_DIR, "intl_dual_momentum_onepager.docx")
    doc.save(path)
    print(f"Saved: {path}")


if __name__ == "__main__":
    build()
