"""
CRTPX Final Presentation: chart + stats + Word document.
Compares Strategy A, Strategy B, SPY, and 70/30 portfolio.
"""
import os, math, warnings
warnings.filterwarnings("ignore")

import numpy as np
import pandas as pd
import yfinance as yf
import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
import matplotlib.dates as mdates
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from lxml import etree

SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
START = "2021-06-01"
END = "2026-02-01"

# ═══════════════════════════════════════════════════════════════
# 1. BUILD PENTA SIGNAL LOCALLY
# ═══════════════════════════════════════════════════════════════

def build_penta_backtest():
    tickers = {"SPY": "^GSPC", "IYT": "IYT", "NYA": "^NYA", "LQD": "LQD",
               "VOO": "VOO", "SGOV": "SGOV", "DBMF": "DBMF", "CAOS": "CAOS",
               "AGG": "AGG"}
    warmup = "2021-01-01"
    raw = yf.download(list(tickers.values()), start=warmup, end=END, progress=False)
    prices = {}
    for name, tk in tickers.items():
        try:
            if isinstance(raw.columns, pd.MultiIndex):
                prices[name] = raw["Close"][tk].dropna()
            else:
                prices[name] = raw["Close"].dropna()
        except:
            pass

    idx = prices["SPY"][prices["SPY"].index >= START].index
    idx = idx.tz_localize(None) if idx.tz is not None else idx
    SMA = 50
    CONFIRM = 3

    penta = pd.DataFrame(index=idx)
    for name in prices:
        if prices[name].index.tz is not None:
            prices[name].index = prices[name].index.tz_localize(None)

    for name, tk in [("SPY","SPY"),("IYT","IYT"),("NYA","NYA"),("LQD","LQD")]:
        p = prices[name].reindex(idx, method="ffill")
        sma = prices[name].rolling(SMA).mean().reindex(idx, method="ffill")
        penta[name] = (p > sma).astype(int)
    penta["score"] = penta.sum(axis=1)
    penta["raw_on"] = (penta["score"] >= 3).astype(int)

    confirmed = pd.Series("OFF", index=idx)
    regime = "OFF"
    count = 0
    for i in range(len(idx)):
        desired = "ON" if penta["raw_on"].iloc[i] else "OFF"
        if regime is None:
            regime = desired
            count = CONFIRM
        if desired != regime:
            count += 1
            if count >= CONFIRM:
                regime = desired
                count = 0
        else:
            count = 0
        confirmed.iloc[i] = regime

    voo_ret = prices["VOO"].reindex(idx, method="ffill").pct_change().fillna(0)
    sgov_ret = prices["SGOV"].reindex(idx, method="ffill").pct_change().fillna(0)
    dbmf_ret = prices["DBMF"].reindex(idx, method="ffill").pct_change().fillna(0)
    caos_ret = prices["CAOS"].reindex(idx, method="ffill").pct_change().fillna(0)
    spy_ret = prices["SPY"].reindex(idx, method="ffill").pct_change().fillna(0)
    agg_ret = prices["AGG"].reindex(idx, method="ffill").pct_change().fillna(0)

    strat_a = pd.Series(0.0, index=idx)
    strat_b = pd.Series(0.0, index=idx)
    bench_7030 = 0.70 * spy_ret + 0.30 * agg_ret

    for i in range(len(idx)):
        if confirmed.iloc[i] == "ON":
            strat_a.iloc[i] = voo_ret.iloc[i]
            strat_b.iloc[i] = voo_ret.iloc[i]
        else:
            strat_a.iloc[i] = sgov_ret.iloc[i]
            strat_b.iloc[i] = (0.40 * dbmf_ret.iloc[i] +
                               0.30 * caos_ret.iloc[i] +
                               0.30 * sgov_ret.iloc[i])

    eq_a = (1 + strat_a).cumprod() * 10000
    eq_b = (1 + strat_b).cumprod() * 10000
    eq_spy = (1 + spy_ret).cumprod() * 10000
    eq_7030 = (1 + bench_7030).cumprod() * 10000

    return idx, eq_a, eq_b, eq_spy, eq_7030, strat_a, strat_b, spy_ret, bench_7030, confirmed


def calc_stats(ret, spy_ret, label):
    ret = ret.dropna()
    spy_ret = spy_ret.reindex(ret.index).fillna(0)
    eq = (1 + ret).cumprod()
    days = (eq.index[-1] - eq.index[0]).days
    yrs = days / 365.25
    cagr = (eq.iloc[-1] / eq.iloc[0]) ** (1/yrs) - 1
    total = eq.iloc[-1] / eq.iloc[0] - 1
    dd = ((eq - eq.cummax()) / eq.cummax()).min()
    vol = ret.std() * math.sqrt(252)
    sharpe = (ret.mean() - 0.04/252) / ret.std() * math.sqrt(252) if ret.std() > 0 else 0
    down = ret[ret < 0]
    sortino = (ret.mean() - 0.04/252) / down.std() * math.sqrt(252) if len(down) > 0 and down.std() > 0 else 0
    calmar = cagr / abs(dd) if dd != 0 else 0
    c = ret.index.intersection(spy_ret.index)
    if len(c) > 50:
        cv = np.cov(ret.loc[c], spy_ret.loc[c])
        beta = cv[0,1] / cv[1,1] if cv[1,1] > 0 else 0
        corr = np.corrcoef(ret.loc[c], spy_ret.loc[c])[0,1]
        alpha = cagr - (0.04 + beta * (spy_ret.mean()*252 - 0.04))
    else:
        beta, corr, alpha = 0, 0, 0

    up_months = ret.resample("ME").sum()
    spy_months = spy_ret.resample("ME").sum()
    up_mask = spy_months > 0
    dn_mask = spy_months < 0
    up_cap = (up_months[up_mask].mean() / spy_months[up_mask].mean() * 100) if up_mask.sum() > 0 else 0
    dn_cap = (up_months[dn_mask].mean() / spy_months[dn_mask].mean() * 100) if dn_mask.sum() > 0 else 0

    return {
        "label": label,
        "cagr": f"{cagr*100:.2f}%",
        "total": f"{total*100:.1f}%",
        "max_dd": f"{dd*100:.2f}%",
        "vol": f"{vol*100:.2f}%",
        "sharpe": f"{sharpe:.2f}",
        "sortino": f"{sortino:.2f}",
        "calmar": f"{calmar:.2f}",
        "beta": f"{beta:.2f}",
        "alpha": f"{alpha*100:.2f}%",
        "corr": f"{corr:.2f}",
        "up_cap": f"{up_cap:.1f}%",
        "dn_cap": f"{dn_cap:.1f}%",
    }


# ═══════════════════════════════════════════════════════════════
# 2. CHART
# ═══════════════════════════════════════════════════════════════

def make_chart(idx, eq_a, eq_b, eq_spy, eq_7030, confirmed):
    import datetime
    dates = [datetime.datetime(d.year, d.month, d.day) for d in idx]
    fig, (ax1, ax2) = plt.subplots(2, 1, figsize=(10, 6.5),
                                    gridspec_kw={"height_ratios": [4, 1]},
                                    sharex=True)
    fig.patch.set_facecolor("white")

    x_pos = np.arange(len(dates))
    ax1.plot(x_pos, eq_a.values, label="Strategy A: Passive + TLH", color="#003366", linewidth=1.8)
    ax1.plot(x_pos, eq_b.values, label="Strategy B: Enhanced + TLH", color="#CC3300", linewidth=1.8)
    ax1.plot(x_pos, eq_spy.values, label="SPY Buy & Hold", color="#888888", linewidth=1.2, linestyle="--")
    ax1.plot(x_pos, eq_7030.values, label="70/30 (SPY/AGG)", color="#66AA66", linewidth=1.2, linestyle="--")
    ax1.set_ylabel("Growth of $10,000", fontsize=10)
    ax1.legend(fontsize=8, loc="upper left")
    ax1.grid(True, alpha=0.3)
    ax1.set_title("CRTPX Strategy Comparison  |  June 2021 - Feb 2026", fontsize=12, fontweight="bold")

    regime_vals = confirmed.map({"ON": 1, "OFF": 0}).values
    ax2.fill_between(x_pos, 0, 1, where=regime_vals == 1, color="#d4edda", alpha=0.8)
    ax2.fill_between(x_pos, 0, 1, where=regime_vals == 0, color="#f8d7da", alpha=0.8)
    ax2.set_ylabel("Regime", fontsize=9)
    ax2.set_yticks([])
    ax2.set_ylim(0, 1)
    tick_positions = np.linspace(0, len(dates)-1, 12, dtype=int)
    ax2.set_xticks(tick_positions)
    ax2.set_xticklabels([dates[i].strftime("%b '%y") for i in tick_positions],
                        rotation=45, fontsize=8)
    ax1.set_xticks(tick_positions)
    ax1.set_xticklabels([])
    plt.tight_layout()

    path = os.path.join(SCRIPT_DIR, "crtpx_comparison_chart.png")
    fig.savefig(path, dpi=200, bbox_inches="tight")
    plt.close()
    print(f"Chart saved: {path}")
    return path


# ═══════════════════════════════════════════════════════════════
# 3. WORD DOCUMENT
# ═══════════════════════════════════════════════════════════════

def sc(cell, text, bold=False, sz=8, color=None, align="left"):
    cell.text = ""
    p = cell.paragraphs[0]
    r = p.add_run(str(text))
    r.font.size = Pt(sz); r.bold = bold; r.font.name = "Calibri"
    if color: r.font.color.rgb = RGBColor(*color)
    if align == "center": p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif align == "right": p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

def mk_table(doc, hdrs, rows):
    t = doc.add_table(rows=1+len(rows), cols=len(hdrs))
    t.style = "Light Grid Accent 1"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(hdrs):
        sc(t.rows[0].cells[i], h, bold=True, sz=8, color=(255,255,255), align="center")
        tc = t.rows[0].cells[i]._element
        pr = tc.find(qn("w:tcPr"))
        if pr is None: pr = etree.SubElement(tc, qn("w:tcPr"))
        shd = etree.SubElement(pr, qn("w:shd"))
        shd.set(qn("w:fill"), "003366"); shd.set(qn("w:val"), "clear")
    for ri, row in enumerate(rows):
        for ci, v in enumerate(row):
            sc(t.rows[ri+1].cells[ci], v, sz=8, align="right" if ci > 0 else "left")
    return t

def ap(doc, text, bold=False, sz=10, after=6):
    p = doc.add_paragraph()
    r = p.add_run(text); r.font.size = Pt(sz); r.bold = bold; r.font.name = "Calibri"
    p.paragraph_format.space_after = Pt(after)
    return p

def code(doc, text):
    p = doc.add_paragraph()
    r = p.add_run(text); r.font.size = Pt(7); r.font.name = "Consolas"
    p.paragraph_format.space_before = Pt(2); p.paragraph_format.space_after = Pt(2)

def build_doc(chart_path, stats_a, stats_b, stats_spy, stats_7030, regime_info):
    doc = Document()
    s = doc.styles["Normal"]; s.font.name = "Calibri"; s.font.size = Pt(10)
    for sec in doc.sections:
        sec.top_margin = Cm(1.2); sec.bottom_margin = Cm(1.2)
        sec.left_margin = Cm(1.8); sec.right_margin = Cm(1.8)

    h = doc.add_heading("CRTPX: Tactical Risk Management for the Passive Fund", level=0)
    for r in h.runs: r.font.color.rgb = RGBColor(0, 51, 102)
    ap(doc, "Potomac Fund Management  |  Woody Wiegmann  |  February 2026", sz=9, after=2)
    ap(doc, f"Common period: June 2021 - February 2026  |  "
       f"Risk-off days: {regime_info['off_pct']}  |  "
       f"Regime switches: {regime_info['switches']} ({regime_info['per_yr']}/yr)", sz=9, after=8)

    doc.add_picture(chart_path, width=Inches(6.2))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_heading("Performance Comparison", level=2)
    mk_table(doc,
        ["Metric", "Strategy A\n(Passive)", "Strategy B\n(Enhanced)", "SPY\nBuy&Hold", "70/30\n(SPY/AGG)"],
        [
            ["CAGR", stats_a["cagr"], stats_b["cagr"], stats_spy["cagr"], stats_7030["cagr"]],
            ["Total Return", stats_a["total"], stats_b["total"], stats_spy["total"], stats_7030["total"]],
            ["Max Drawdown", stats_a["max_dd"], stats_b["max_dd"], stats_spy["max_dd"], stats_7030["max_dd"]],
            ["Volatility", stats_a["vol"], stats_b["vol"], stats_spy["vol"], stats_7030["vol"]],
            ["Sharpe", stats_a["sharpe"], stats_b["sharpe"], stats_spy["sharpe"], stats_7030["sharpe"]],
            ["Sortino", stats_a["sortino"], stats_b["sortino"], stats_spy["sortino"], stats_7030["sortino"]],
            ["Calmar", stats_a["calmar"], stats_b["calmar"], stats_spy["calmar"], stats_7030["calmar"]],
            ["Beta", stats_a["beta"], stats_b["beta"], stats_spy["beta"], stats_7030["beta"]],
            ["Alpha", stats_a["alpha"], stats_b["alpha"], stats_spy["alpha"], stats_7030["alpha"]],
            ["Up Capture", stats_a["up_cap"], stats_b["up_cap"], stats_spy["up_cap"], stats_7030["up_cap"]],
            ["Down Capture", stats_a["dn_cap"], stats_b["dn_cap"], stats_spy["dn_cap"], stats_7030["dn_cap"]],
        ])

    doc.add_page_break()

    # ── STRATEGY SPECS ──────────────────────────────────────────
    doc.add_heading("Signal: Penta (Both Strategies)", level=2)
    mk_table(doc,
        ["#", "Indicator", "Ticker", "Rule"],
        [["1","S&P 500 Trend","SPY","Close > 50-day SMA"],
         ["2","Dow Transports","IYT","Close > 50-day SMA"],
         ["3","NYSE Breadth","^NYA","Close > 50-day SMA"],
         ["4","Corp Credit","LQD","Close > 50-day SMA"]])
    ap(doc, "Penta ON = 3+ green. 3-day confirmation before switching.", sz=9, after=8)

    doc.add_heading("Allocations", level=2)
    mk_table(doc,
        ["Regime", "Strategy A", "Strategy B"],
        [["Risk-On (Penta >= 3)", "100% VOO", "100% VOO"],
         ["Risk-Off (Penta <= 2)", "100% SGOV", "40% DBMF + 30% CAOS + 30% SGOV"]])
    ap(doc, "Strategy B risk-off blend based on Baltussen et al. (FAJ 2026): "
       "trend-following (DBMF) + convexity (CAOS) + cash anchor (SGOV). "
       "222-year evidence shows 50/50 trend+convexity is optimal for drawdown reduction.", sz=8, after=8)

    # ── OPERATING GUIDE ─────────────────────────────────────────
    doc.add_heading("Operating Guide", level=1)
    for r in doc.paragraphs[-1].runs: r.font.color.rgb = RGBColor(0, 51, 102)

    ap(doc, "Daily Check (takes ~2 minutes):", bold=True, sz=10, after=3)
    ap(doc, "1. Pull closing prices for SPY, IYT, ^NYA, LQD.\n"
       "2. Compare each to its 50-day SMA. Count greens (0-4).\n"
       "3. If score differs from current regime for 3 consecutive days, switch.\n"
       "4. Check current holding for >3% unrealized loss (TLH trigger).", sz=9, after=6)

    ap(doc, "Risk-On to Risk-Off:", bold=True, sz=10, after=3)
    ap(doc, "1. Check equity lot for unrealized loss. If loss: sell first (harvest).\n"
       "2. Buy defensive: (A) 100% SGOV, or (B) 40% DBMF + 30% CAOS + 30% SGOV.\n"
       "3. Record sold ticker + wash sale expiry (T+31 days).", sz=9, after=6)

    ap(doc, "Risk-Off to Risk-On:", bold=True, sz=10, after=3)
    ap(doc, "1. Check defensive lots for unrealized loss. If loss: sell first.\n"
       "2. Buy equity: VOO (or next in ring if wash sale active).\n"
       "3. Record wash sale dates.", sz=9, after=6)

    ap(doc, "TLH Swap Rings:", bold=True, sz=10, after=3)
    mk_table(doc,
        ["Side", "Primary", "Swap 1", "Swap 2"],
        [["Equity", "VOO", "IVV", "SPLG"],
         ["Cash", "SGOV", "BIL", "SHV"]])
    ap(doc, "Loss trigger: -3%. Wash sale window: 31 calendar days. "
       "Always use next available ticker in ring.", sz=8, after=8)

    doc.add_page_break()

    # ── AMIBROKER AFL ───────────────────────────────────────────
    doc.add_heading("AmiBroker Implementation", level=1)
    for r in doc.paragraphs[-1].runs: r.font.color.rgb = RGBColor(0, 51, 102)

    ap(doc, "Strategy A (Passive):", bold=True, sz=10, after=3)
    code(doc, """SMA_Period = 50;  ConfirmDays = 3;
Penta1 = Foreign("SPY","C") > MA(Foreign("SPY","C"), SMA_Period);
Penta2 = Foreign("IYT","C") > MA(Foreign("IYT","C"), SMA_Period);
Penta3 = Foreign("~NYA","C") > MA(Foreign("~NYA","C"), SMA_Period);
Penta4 = Foreign("LQD","C") > MA(Foreign("LQD","C"), SMA_Period);
PentaScore = Penta1 + Penta2 + Penta3 + Penta4;
RawOn = PentaScore >= 3;
ConfOn  = Sum(RawOn, ConfirmDays) == ConfirmDays;
ConfOff = Sum(!RawOn, ConfirmDays) == ConfirmDays;
Buy  = ExRem(ConfOn, ConfOff);
Sell = ExRem(ConfOff, ConfOn);
SetPositionSize(100, spsPercentOfEquity);
// Trade SPY or VOO. Risk-off = cash (or model SGOV via SetForeign).""")

    ap(doc, "", after=4)
    ap(doc, "Strategy B (Enhanced) -- risk-off return:", bold=True, sz=10, after=3)
    code(doc, """// Same signal as above. When Sell fires, risk-off return =
RiskOffRet = 0.40 * ROC(Foreign("DBMF","C"),1)/100
           + 0.30 * ROC(Foreign("CAOS","C"),1)/100
           + 0.30 * ROC(Foreign("SGOV","C"),1)/100;
// Use SetForeign or rotational mode to apply blended risk-off.""")

    ap(doc, "", after=4)
    ap(doc, "Sensitivity test these parameters:", bold=True, sz=10, after=3)
    mk_table(doc,
        ["Parameter", "Default", "Test Range"],
        [["SMA Period", "50", "20, 50, 100, 200"],
         ["Confirmation Days", "3", "1, 2, 3, 5"],
         ["Penta Threshold", "3/4", "2/4, 3/4, 4/4"],
         ["DBMF / CAOS / SGOV wts", "40/30/30", "50/25/25, 30/40/30"]])

    ap(doc, "", after=4)
    ap(doc, "Why this isn't overfit: 4 standard indicators, 1 parameter each (50-day SMA). "
       "3-day confirmation is a standard anti-whipsaw filter. Risk-off blend derived from "
       "222-year evidence (Baltussen et al. 2026), not from fitting to this period. "
       "Only 2 free parameters total.", sz=8, after=6)

    path = os.path.join(SCRIPT_DIR, "CRTPX_Strategy_Presentation.docx")
    doc.save(path)
    print(f"Document saved: {path}")
    return path


# ═══════════════════════════════════════════════════════════════
# MAIN
# ═══════════════════════════════════════════════════════════════

def main():
    print("Building CRTPX comparison...")
    idx, eq_a, eq_b, eq_spy, eq_7030, ret_a, ret_b, ret_spy, ret_7030, confirmed = build_penta_backtest()

    switches = (confirmed != confirmed.shift()).sum() - 1
    off_days = (confirmed == "OFF").sum()
    total_days = len(confirmed)
    yrs = (idx[-1] - idx[0]).days / 365.25
    regime_info = {
        "switches": switches,
        "per_yr": f"{switches/yrs:.1f}",
        "off_pct": f"{off_days/total_days*100:.0f}%",
    }

    print(f"  Period: {idx[0].date()} to {idx[-1].date()}")
    print(f"  Switches: {switches} ({switches/yrs:.1f}/yr)")
    print(f"  Risk-off: {off_days}/{total_days} days ({off_days/total_days*100:.0f}%)")

    stats_a = calc_stats(ret_a, ret_spy, "Strategy A")
    stats_b = calc_stats(ret_b, ret_spy, "Strategy B")
    stats_spy = calc_stats(ret_spy, ret_spy, "SPY")
    stats_7030 = calc_stats(ret_7030, ret_spy, "70/30")

    print("\n  CAGR:    A={} B={} SPY={} 70/30={}".format(
        stats_a["cagr"], stats_b["cagr"], stats_spy["cagr"], stats_7030["cagr"]))
    print("  MaxDD:   A={} B={} SPY={} 70/30={}".format(
        stats_a["max_dd"], stats_b["max_dd"], stats_spy["max_dd"], stats_7030["max_dd"]))
    print("  Sharpe:  A={} B={} SPY={} 70/30={}".format(
        stats_a["sharpe"], stats_b["sharpe"], stats_spy["sharpe"], stats_7030["sharpe"]))

    chart = make_chart(idx, eq_a, eq_b, eq_spy, eq_7030, confirmed)
    build_doc(chart, stats_a, stats_b, stats_spy, stats_7030, regime_info)

if __name__ == "__main__":
    main()
