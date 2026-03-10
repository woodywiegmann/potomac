"""
Generate one-pager Word doc: 40-Ticker International Dual Momentum Universe
"""

import sys
try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn
except ImportError:
    sys.exit("Required: pip install python-docx")

UNIVERSE = {
    "Dev Country": [
        ("EWJ",  "Japan",         "$20.2B", "1996", "Core DM; 2nd largest equity market globally, low corr to Europe"),
        ("EWG",  "Germany",       "$1.8B",  "1996", "Europe's largest economy; industrials/auto exposure"),
        ("EWQ",  "France",        "$386M",  "1996", "Luxury, energy, defense tilt; pairs with EWG for Eurozone breadth"),
        ("EWI",  "Italy",         "$681M",  "1996", "Banks/energy heavy; distinct cycle vs Northern Europe"),
        ("EWD",  "Sweden",        "$360M",  "1996", "High-beta Nordic; tech & industrials tilt"),
        ("EWL",  "Switzerland",   "$1.7B",  "1996", "Defensive quality; pharma/consumer staples overweight"),
        ("EWP",  "Spain",         "$924M",  "1996", "Financials/utilities heavy; distinct from core Eurozone"),
        ("EWH",  "Hong Kong",     "$770M",  "1996", "Gateway to China; lowest corr among dev country ETFs (0.43)"),
        ("EWS",  "Singapore",     "$794M",  "1996", "Financials/REITs hub; Asia-Pacific diversifier"),
        ("EDEN", "Denmark",       "$237M",  "2012", "Healthcare-dominated (Novo Nordisk); unique return profile"),
    ],
    "Dev Factor": [
        ("IHDG", "Intl Hedged Qual Div Growth", "$2.4B", "2014",
         "Currency-hedged quality dividend growers; decorrelates FX risk from equity returns"),
    ],
    "Dev Thematic": [
        ("RING", "Global Gold Miners",       "$4.0B",  "2012", "Gold beta; lowest avg corr among thematics (0.37)"),
        ("SIL",  "Silver Miners",            "$7.3B",  "2010", "Precious metals + industrial demand dual driver"),
        ("URA",  "Uranium",                  "$7.6B",  "2010", "Nuclear renaissance theme; low corr to broad equity (0.41)"),
        ("KXI",  "Global Consumer Staples",  "$1.0B",  "2006", "Defensive sector rotation; anti-cyclical diversifier"),
        ("LIT",  "Lithium & Battery Tech",   "$1.8B",  "2010", "EV supply chain; distinct from base metals cycle"),
        ("REMX", "Rare Earth & Strat Metals","$3.0B",  "2010", "Strategic materials; China supply-chain exposure"),
        ("COPX", "Copper Miners",            "$8.3B",  "2010", "Infrastructure/electrification pure-play"),
        ("PICK", "Global Metals & Mining",   "$2.1B",  "2012", "Broad base metals diversifier across geographies"),
        ("GNR",  "S&P Global Nat Resources", "$4.8B",  "2010", "Energy + agriculture + metals blend"),
        ("CGW",  "Global Water",             "$1.1B",  "2007", "Secular theme; utilities-like defensiveness"),
        ("GII",  "Global Infrastructure",    "$816M",  "2007", "Utilities/transport/energy infra; low equity beta"),
        ("INFL", "Inflation Beneficiaries",  "$1.5B",  "2021", "Real-asset tilt; hedges inflationary regimes"),
        ("MOO",  "Agribusiness",             "$1.0B",  "2007", "Food supply chain; weather/commodity cycle exposure"),
    ],
    "EM Country": [
        ("EWT",  "Taiwan",         "$8.8B",  "2000", "Semiconductor-heavy; TSMC concentration = unique risk/reward"),
        ("EWZ",  "Brazil",         "$9.7B",  "2000", "Commodity-driven EM; distinct LatAm cycle"),
        ("INDA", "India",          "$9.3B",  "2012", "Fastest-growing large economy; low corr to China (0.28)"),
        ("FXI",  "China Large-Cap","$6.2B",  "2004", "Largest EM market; policy-driven return stream"),
        ("EWY",  "South Korea",    "$18.0B", "2000", "Tech/memory chip cycle; high beta to global trade"),
        ("EWW",  "Mexico",         "$2.6B",  "1996", "Nearshoring beneficiary; peso carry trade exposure"),
        ("ILF",  "Latin America 40","$4.4B", "2001", "LatAm breadth beyond Brazil (Chile, Colombia, Peru)"),
        ("ECH",  "Chile",          "$1.3B",  "2007", "Copper/lithium economy; EM commodity pure-play"),
        ("TUR",  "Turkey",         "$350M",  "2008", "THE diversifier -- lowest avg corr in universe (0.18)"),
        ("ARGT", "Argentina",      "$806M",  "2011", "Frontier-like volatility; reform cycle optionality"),
        ("VNM",  "Vietnam",        "$655M",  "2009", "Frontier-EM; manufacturing relocation beneficiary (0.26 avg corr)"),
        ("THD",  "Thailand",       "$347M",  "2008", "ASEAN exposure; tourism + manufacturing economy"),
        ("EWM",  "Malaysia",       "$388M",  "1996", "Commodity + tech blend; ASEAN diversifier"),
        ("EIDO", "Indonesia",      "$311M",  "2010", "Largest ASEAN economy; demographics-driven growth"),
    ],
    "EM Broad": [
        ("KSA",  "Saudi Arabia",   "$662M",  "2015", "Vision 2030 diversification; 0.33 avg corr (near lowest)"),
        ("KWEB", "China Internet", "$6.9B",  "2013", "Tech/consumer platform plays; distinct from FXI (SOE-heavy)"),
    ],
}

STATS = {
    "Total tickers": "40",
    "Developed": "24 (60%)",
    "Emerging Markets": "16 (40%)",
    "Avg pairwise correlation (full 71)": "0.5632",
    "Avg pairwise correlation (trimmed 40)": "0.4750",
    "Correlation reduction": "-15.7%",
    "Pairs > 0.90 (full 71)": "65",
    "Pairs > 0.90 (trimmed 40)": "7",
    "Lookback": "5 years daily returns",
}


def set_cell_shading(cell, color):
    shading = cell._element.get_or_add_tcPr()
    sh = shading.makeelement(qn("w:shd"), {
        qn("w:fill"): color, qn("w:val"): "clear"})
    shading.append(sh)


def add_styled_table(doc, headers, rows, col_widths=None, header_color="1F4E79"):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"

    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(8)
                run.font.color.rgb = RGBColor(255, 255, 255)
        set_cell_shading(cell, header_color)

    for r_idx, row_data in enumerate(rows):
        for c_idx, val in enumerate(row_data):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = str(val)
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(7.5)
        if r_idx % 2 == 1:
            for c_idx in range(len(headers)):
                set_cell_shading(table.rows[r_idx + 1].cells[c_idx], "EBF5FB")

    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Inches(w)

    return table


def main():
    doc = Document()

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(9)
    style.paragraph_format.space_after = Pt(2)
    style.paragraph_format.space_before = Pt(1)

    for section in doc.sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.4)
        section.left_margin = Inches(0.5)
        section.right_margin = Inches(0.5)

    # Title
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("International Dual Momentum Strategy\nProposed 40-ETF Universe")
    run.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(31, 78, 121)

    # Subtitle
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = sub.add_run("Potomac Fund Management  |  March 2026")
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(100, 100, 100)

    # Overview
    doc.add_paragraph()
    h = doc.add_paragraph()
    run = h.add_run("Overview")
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(31, 78, 121)

    overview_text = (
        "This universe was constructed for a CRTOX-style international dual momentum strategy. "
        "Starting from a 71-ETF candidate pool (all international, >=$200M AUM), we applied a "
        "two-phase optimization: (1) hard de-duplication of near-identical pairs (>0.95 correlation), "
        "then (2) constrained greedy trimming to minimize average pairwise correlation while "
        "maintaining at least 60% developed-market representation. The result is a 40-ETF universe "
        "with average pairwise correlation of 0.4750 (vs. 0.5632 for the full candidate set), "
        "providing maximum signal differentiation for blended-lookback momentum ranking."
    )
    p = doc.add_paragraph(overview_text)
    p.paragraph_format.space_after = Pt(4)

    # Strategy parameters
    h = doc.add_paragraph()
    run = h.add_run("Strategy Parameters")
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(31, 78, 121)

    params = [
        ("Signal", "Blended momentum: average of 1-month, 3-month, 6-month, and 12-month trailing returns"),
        ("Ranking", "All 40 ETFs ranked by composite score each month"),
        ("Holdings", "Top 7, equal-weight (~14.3% each)"),
        ("Absolute Momentum", "Composite score must be > 0 to be held; failing slots go to cash"),
        ("Risk-Off", "BIL (SPDR Bloomberg 1-3 Month T-Bill ETF)"),
        ("Rebalance", "Monthly (last trading day)"),
    ]
    add_styled_table(doc, ["Parameter", "Detail"],
                     params, col_widths=[1.2, 5.8], header_color="1F4E79")

    # Correlation stats
    doc.add_paragraph()
    h = doc.add_paragraph()
    run = h.add_run("Correlation Summary")
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(31, 78, 121)

    stat_rows = list(STATS.items())
    add_styled_table(doc, ["Metric", "Value"], stat_rows,
                     col_widths=[3.0, 4.0], header_color="1F4E79")

    # Universe tables by bucket
    doc.add_page_break()

    title2 = doc.add_paragraph()
    title2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title2.add_run("Universe Detail: 40 ETFs")
    run.bold = True
    run.font.size = Pt(13)
    run.font.color.rgb = RGBColor(31, 78, 121)

    for bucket_name, tickers in UNIVERSE.items():
        doc.add_paragraph()
        h = doc.add_paragraph()
        run = h.add_run(f"{bucket_name} ({len(tickers)})")
        run.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(31, 78, 121)

        rows = [(t[0], t[1], t[2], t[3], t[4]) for t in tickers]
        add_styled_table(
            doc,
            ["Ticker", "Name", "AUM", "Incep.", "Rationale"],
            rows,
            col_widths=[0.5, 1.5, 0.6, 0.5, 3.9],
            header_color="2C6E49" if "EM" in bucket_name else "1F4E79"
        )

    # Dropped tickers summary
    doc.add_paragraph()
    h = doc.add_paragraph()
    run = h.add_run("Phase 1 De-Duplication (11 Tickers Removed)")
    run.bold = True
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(31, 78, 121)

    dedup_rows = [
        ("FLJP",  "EWJ",  "0.996", "Same Japan exposure"),
        ("FLKR",  "EWY",  "0.992", "Same South Korea"),
        ("FLBR",  "EWZ",  "0.990", "Same Brazil"),
        ("FLGB",  "EWU",  "0.989", "Same United Kingdom"),
        ("FLIN",  "INDA", "0.982", "Same India"),
        ("EPI",   "INDA", "0.964", "India earnings tilt, near-identical"),
        ("URNM",  "URA",  "0.961", "Same uranium miners"),
        ("SILJ",  "SIL",  "0.971", "Junior silver ~ silver miners"),
        ("GUNR",  "GNR",  "0.984", "Same natural resources"),
        ("CQQQ",  "KWEB", "0.926", "Both China tech/internet"),
        ("EEMA",  "AVEM", "0.959", "EM Asia subset of EM broad"),
    ]
    add_styled_table(
        doc,
        ["Dropped", "Kept", "Corr", "Reason"],
        dedup_rows,
        col_widths=[0.6, 0.6, 0.5, 5.3],
        header_color="922B21"
    )

    doc.add_paragraph()
    h = doc.add_paragraph()
    run = h.add_run("Phase 2 Greedy Trim (20 Tickers Removed, Dev >= 60% Floor)")
    run.bold = True
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(31, 78, 121)

    trim_rows = [
        ("AVDE",  "Dev Factor",  "0.724", "Highest avg corr; broad basket overlaps everything"),
        ("PXF",   "Dev Factor",  "0.711", "RAFI dev ex-US; near-identical to AVDE"),
        ("AVDV",  "Dev Factor",  "0.701", "Intl small-cap value; 0.97 corr with SCZ, DLS"),
        ("SCZ",   "Dev Factor",  "0.689", "EAFE small-cap; overlaps AVDV, DLS"),
        ("DLS",   "Dev Factor",  "0.678", "Intl small-cap div; 0.98 corr with SCZ"),
        ("MXI",   "Dev Thematic","0.682", "Global materials; subsumed by GNR, PICK, COPX"),
        ("IVLU",  "Dev Factor",  "0.659", "Intl value factor; 0.97 with PXF, AVDE"),
        ("IMTM",  "Dev Factor",  "0.664", "Intl momentum factor; 0.94 with AVDE"),
        ("EMXC",  "EM Broad",   "0.663", "EM ex-China; 0.92 with AVEM"),
        ("EWA",   "Dev Country", "0.644", "Australia; high corr to EWC, materials complex"),
        ("AVEM",  "EM Broad",   "0.649", "Avantis EM; broad basket, overlaps country ETFs"),
        ("HDEF",  "Dev Factor",  "0.640", "EAFE high div; 0.92 with PXF, IVLU"),
        ("EWC",   "Dev Country", "0.634", "Canada; resources/energy overlap with commodities"),
        ("IVAL",  "Dev Factor",  "0.626", "Intl quant value; 0.91 with AVDV, IVLU"),
        ("DEM",   "EM Broad",   "0.621", "EM high div; 0.90 with AVEM"),
        ("IPAC",  "Dev Factor",  "0.616", "Core Pacific; 0.96 with EWJ"),
        ("EWU",   "Dev Country", "0.612", "United Kingdom; dropped at dev floor constraint"),
        ("GVAL",  "EM Broad",   "0.577", "Global value; broad basket"),
        ("EWN",   "Dev Country", "0.605", "Netherlands; high corr to EWG, EWQ"),
        ("GRID",  "Dev Thematic","0.591", "Smart grid/infra; overlaps GII, CGW"),
    ]
    add_styled_table(
        doc,
        ["Ticker", "Bucket", "Avg Corr", "Reason for Removal"],
        trim_rows,
        col_widths=[0.6, 1.0, 0.6, 4.8],
        header_color="922B21"
    )

    out = r"c:\Users\WoodyWiegmann\OneDrive - PFM\Desktop\Potomac\Intl_DualMom_40ETF_Universe.docx"
    doc.save(out)
    print(f"Saved: {out}")


if __name__ == "__main__":
    main()
