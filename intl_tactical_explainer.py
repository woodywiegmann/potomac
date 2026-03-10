"""
Generate 2-page strategy explanation document for the International Tactical
Fund with Breadth/Trend Overlay (60% threshold variant).
"""

import sys
try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn
except ImportError:
    sys.exit("Required: pip install python-docx")


def set_cell_shading(cell, color):
    shading = cell._element.get_or_add_tcPr()
    sh = shading.makeelement(qn("w:shd"), {qn("w:fill"): color, qn("w:val"): "clear"})
    shading.append(sh)


def styled_table(doc, headers, rows, header_color="1F4E79", col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"

    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(8.5)
                run.font.color.rgb = RGBColor(255, 255, 255)
        set_cell_shading(cell, header_color)

    for r_idx, row_data in enumerate(rows):
        for c_idx, val in enumerate(row_data):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = str(val)
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(8.5)
        if r_idx % 2 == 1:
            for c_idx in range(len(headers)):
                set_cell_shading(table.rows[r_idx + 1].cells[c_idx], "EBF5FB")

    if col_widths:
        for row in table.rows:
            for i, w in enumerate(col_widths):
                row.cells[i].width = Inches(w)

    return table


def add_heading(doc, text, size=12, color="1F4E79"):
    r, g, b = int(color[0:2], 16), int(color[2:4], 16), int(color[4:6], 16)
    h = doc.add_paragraph()
    run = h.add_run(text)
    run.bold = True
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor(r, g, b)
    h.paragraph_format.space_before = Pt(4)
    h.paragraph_format.space_after = Pt(2)
    return h


def add_body(doc, text, size=9.5, bold=False, italic=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(3)
    return p


def main():
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(9.5)
    for section in doc.sections:
        section.top_margin = Inches(0.55)
        section.bottom_margin = Inches(0.45)
        section.left_margin = Inches(0.65)
        section.right_margin = Inches(0.65)

    # Title
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("International Tactical Fund")
    run.bold = True
    run.font.size = Pt(18)
    run.font.color.rgb = RGBColor(31, 78, 121)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = sub.add_run("How the Strategy Makes Trades")
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(80, 80, 80)
    sub.paragraph_format.space_after = Pt(6)

    # ---------------------------------------------------------------
    # SECTION 1: OVERVIEW
    # ---------------------------------------------------------------
    add_heading(doc, "What This Strategy Does")

    add_body(doc,
        "This is a rules-based international equity strategy. Every month, it asks "
        "two questions: (1) Is the global environment healthy enough to be invested? "
        "and (2) Which international markets have the strongest momentum right now? "
        "Based on the answers, it either owns the top-performing international ETFs "
        "or steps aside into short-term Treasuries."
    )

    # ---------------------------------------------------------------
    # SECTION 2: THE INVESTMENT UNIVERSE
    # ---------------------------------------------------------------
    add_heading(doc, "Step 1: The Investment Universe")

    add_body(doc,
        "The strategy selects from a universe of 40 international ETFs. These were "
        "chosen to provide broad geographic coverage while minimizing overlap. "
        "The universe is roughly 60% developed markets and 40% emerging markets, "
        "spanning single-country funds, thematic/commodity exposure, and regional baskets."
    )

    universe = [
        ("Developed Country (10)", "EWJ (Japan), EWG (Germany), EWQ (France), EWI (Italy), "
         "EWD (Sweden), EWL (Switzerland), EWP (Spain), EWH (Hong Kong), EWS (Singapore), EDEN (Denmark)"),
        ("Developed Factor (1)", "IHDG (Intl Hedged Quality Dividend Growth)"),
        ("Developed Thematic (13)", "RING, SIL, URA, KXI, LIT, REMX, COPX, PICK, GNR, CGW, GII, INFL, MOO"),
        ("EM Country (14)", "EWT, EWZ, INDA, FXI, EWY, EWW, ILF, ECH, TUR, ARGT, VNM, THD, EWM, EIDO"),
        ("EM Broad (2)", "KSA (Saudi Arabia), KWEB (China Internet)"),
    ]
    styled_table(doc, ["Category", "ETFs"], universe, header_color="2C6E49",
                 col_widths=[1.8, 5.2])

    # ---------------------------------------------------------------
    # SECTION 3: THE MACRO OVERLAY
    # ---------------------------------------------------------------
    add_heading(doc, "Step 2: Check the Macro Environment (Monthly)")

    add_body(doc,
        "Before ranking any ETFs, the strategy checks two independent macro signals "
        "to determine whether it is safe to be invested at all."
    )

    add_body(doc, "Signal A -- Global Breadth", bold=True)
    add_body(doc,
        "Count how many of the 31 MSCI single-country ETFs (iShares) are trading above "
        "their 200-day simple moving average. If 60% or more are above, breadth is positive. "
        "This measures how widespread the global equity rally is -- when most countries are "
        "in uptrends, the environment favors risk assets."
    )

    add_body(doc, "Signal B -- Ex-US Market Trend", bold=True)
    add_body(doc,
        "Check whether ACWX (iShares MSCI ACWI ex-US) is trading above its own 200-day "
        "simple moving average. This is a single, broad read on whether international "
        "equities as a whole are in an uptrend."
    )

    add_body(doc, "The Decision Rule", bold=True)
    add_body(doc,
        "If EITHER signal is positive, the strategy stays invested and proceeds "
        "to the momentum ranking step below. It goes to 100% cash (BIL -- SPDR "
        "Bloomberg 1-3 Month T-Bill ETF) ONLY if BOTH signals are negative. "
        "This \"OR gate\" design avoids whipsaws: a single positive signal is "
        "enough to stay in the market."
    )

    scenarios = [
        ("Breadth >= 60% AND ACWX > 200d SMA", "INVESTED", "Both signals confirm -- full risk-on"),
        ("Breadth >= 60% BUT ACWX < 200d SMA", "INVESTED", "Breadth alone keeps the portfolio invested"),
        ("Breadth < 60% BUT ACWX > 200d SMA", "INVESTED", "Trend alone keeps the portfolio invested"),
        ("Breadth < 60% AND ACWX < 200d SMA", "100% CASH (BIL)", "Both signals negative -- exit to T-Bills"),
    ]
    styled_table(doc, ["Condition", "Action", "Rationale"], scenarios,
                 header_color="8B0000", col_widths=[2.8, 1.5, 2.7])

    # ---------------------------------------------------------------
    # SECTION 2b: COMPOSITE RISK-ON/RISK-OFF (ALTERNATIVE)
    # ---------------------------------------------------------------
    add_heading(doc, "Step 2 (Alternative): Composite Risk-On/Risk-Off")

    add_body(doc,
        "An alternative to the two-signal overlay above is a composite of multiple "
        "risk-on/risk-off signals, designed to maximize Calmar while staying invested "
        "60–90% of the time. The composite includes: (1) Breadth (% MSCI country ETFs "
        "above 200d SMA); (2) ACWX trend (above 200d SMA); (3) ACWX blended momentum "
        "(1/3/6/12m, normalized); (4) Volatility (VIX-based: low vol = risk-on); "
        "(5) Credit (BNDX above 200d SMA); (6) Relative strength (ACWX vs SPY 12m); "
        "(7) RSI(5) on ACWX (early trend, 50% equilibrium); (8) WMA/IWMA trend "
        "(WMA > IWMA = trend); (9) Turtle Donchian (20/55-day breakout trend)."
    )
    add_body(doc,
        "The composite score (0–1) is a weighted sum of these signals. Equity exposure "
        "is graduated: equity weight = max(composite, floor), with a floor (e.g. 25%) "
        "so the portfolio is rarely 0% invested. Run intl_composite_risk_backtest.py "
        "for backtests and intl_risk_dashboard.py for a monthly HTML dashboard."
    )

    # ---------------------------------------------------------------
    # SECTION 4: MOMENTUM RANKING
    # ---------------------------------------------------------------
    add_heading(doc, "Step 3: Rank by Blended Momentum")

    add_body(doc,
        "If the macro overlay is positive, the strategy scores all 40 ETFs using a "
        "blended momentum signal. For each ETF, it calculates four trailing total returns:"
    )

    lookbacks = [
        ("1-month return", "~21 trading days", "Captures recent acceleration"),
        ("3-month return", "~63 trading days", "Short-term trend"),
        ("6-month return", "~126 trading days", "Medium-term trend"),
        ("12-month return", "~252 trading days", "Long-term trend"),
    ]
    styled_table(doc, ["Lookback", "Window", "Purpose"], lookbacks,
                 col_widths=[1.5, 1.5, 4.0])

    add_body(doc,
        "The composite momentum score is the simple average of these four returns. "
        "This blended approach avoids over-reliance on any single lookback period and "
        "smooths out noise."
    )

    # ---------------------------------------------------------------
    # SECTION 5: PORTFOLIO CONSTRUCTION
    # ---------------------------------------------------------------
    add_heading(doc, "Step 4: Build the Portfolio")

    add_body(doc,
        "The 40 ETFs are ranked by composite momentum score from highest to lowest. "
        "The strategy then applies two filters:"
    )

    add_body(doc,
        "1. Relative momentum: Select the top 7 ETFs by score.\n"
        "2. Absolute momentum: Each of those 7 must have a positive composite score "
        "(i.e., its blended return must be greater than zero). Any ETF with a negative "
        "score is replaced with BIL."
    )

    add_body(doc,
        "The 7 slots are equal-weighted at approximately 14.3% each. If, say, only 5 of "
        "the top 7 have positive scores, the portfolio holds those 5 ETFs at 14.3% each "
        "and allocates the remaining 2 slots (28.6%) to BIL."
    )

    add_body(doc, "Example month:", bold=True, italic=True)
    add_body(doc,
        "Top 7 by score: INDA (+8.2%), TUR (+7.1%), ARGT (+5.9%), COPX (+4.3%), "
        "EWZ (+2.1%), EWJ (+0.4%), SIL (-1.2%). SIL has a negative composite score, "
        "so the portfolio holds 6 ETFs at 14.3% each and BIL at 14.3%. Total equity "
        "exposure: ~85.7%."
    )

    # ---------------------------------------------------------------
    # SECTION 6: REBALANCING
    # ---------------------------------------------------------------
    add_heading(doc, "Step 5: Rebalance")

    add_body(doc,
        "The entire process runs once per month on the last trading day, 30 minutes "
        "before market close. The strategy re-checks both macro signals, re-scores "
        "all 40 ETFs, and adjusts holdings accordingly. Positions are only traded if "
        "the weight difference exceeds 2% to minimize unnecessary turnover."
    )

    # ---------------------------------------------------------------
    # SECTION 7: SUMMARY FLOW
    # ---------------------------------------------------------------
    add_heading(doc, "Decision Flow Summary")

    add_body(doc,
        "Each month, the logic follows this sequence:\n\n"
        "  1. Calculate global breadth (% of 31 MSCI country ETFs above 200d SMA)\n"
        "  2. Check ACWX trend (above/below its 200d SMA)\n"
        "  3. If BOTH are negative --> sell everything, hold 100% BIL\n"
        "  4. If EITHER is positive --> proceed to momentum ranking\n"
        "  5. Score all 40 ETFs using blended momentum (avg of 1m/3m/6m/12m returns)\n"
        "  6. Select top 7 by score\n"
        "  7. Drop any with negative composite score; replace with BIL\n"
        "  8. Equal-weight the 7 slots (~14.3% each)\n"
        "  9. Hold until next month-end rebalance"
    )

    # ---------------------------------------------------------------
    # SECTION 8: BACKTEST STATS
    # ---------------------------------------------------------------
    add_heading(doc, "Backtest Results (Jan 2016 - Feb 2026)", size=11)

    stats = [
        ("CAGR", "12.2%"),
        ("Total Return", "222.4%"),
        ("Max Drawdown", "21.8%"),
        ("Sharpe Ratio", "0.511"),
        ("Sortino Ratio", "0.450"),
        ("Alpha (vs EFA)", "0.049"),
        ("Beta (vs EFA)", "0.337"),
        ("Annualized Volatility", "12.6%"),
        ("Win Rate", "63%"),
        ("Profit/Loss Ratio", "1.22"),
    ]
    styled_table(doc, ["Metric", "Value"], stats, col_widths=[2.5, 1.5])

    p = doc.add_paragraph()
    run = p.add_run(
        "Potomac Fund Management  |  Strategy backtest via QuantConnect LEAN"
    )
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(150, 150, 150)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    out = r"c:\Users\WoodyWiegmann\OneDrive - PFM\Desktop\Potomac\Intl_Tactical_How_Trades_Are_Made.docx"
    doc.save(out)
    print(f"Saved: {out}")


if __name__ == "__main__":
    main()
