"""
CRTOX Universe Expansion One-Pager
===================================
Fetches live price data, computes pairwise correlations for current vs
proposed universe, and generates a clean Word document.
"""

import math, os, sys, warnings
warnings.filterwarnings("ignore")

try:
    import yfinance as yf
    import pandas as pd
    import numpy as np
except ImportError:
    print("Required: pip install yfinance pandas numpy")
    sys.exit(1)

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml

SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
OUT_DIR = r"c:\Users\WoodyWiegmann\OneDrive - PFM\Desktop\Portfolio Ideas"

NAVY = RGBColor(0x1A, 0x3C, 0x5E)
DARK = RGBColor(0x33, 0x33, 0x33)
GRAY = RGBColor(0x66, 0x66, 0x66)
RED  = RGBColor(0xC0, 0x39, 0x2B)
GREEN = RGBColor(0x1E, 0x7D, 0x32)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)

# ═══════════════════════════════════════════════════════════════════════════════
# UNIVERSES
# ═══════════════════════════════════════════════════════════════════════════════

CURRENT_UNIVERSE = {
    "SIL":   "Silver Miners",
    "SMH":   "Semiconductors",
    "XME":   "Metals & Mining",
    "ITA":   "Aerospace & Defense",
    "IBB":   "Biotech",
    "ARKK":  "Innovation / Growth",
    "ILF":   "Latin America",
    "EFV":   "Intl Value (EAFE)",
    "IWO":   "Small Cap Growth",
    "SDIV":  "Global SuperDividend",
    "IGV":   "Software",
    "IAI":   "Broker-Dealers",
    "SOXX":  "Semiconductors (alt)",
}

PROPOSED_UNIVERSE = {
    "SIL":   "Silver Miners",
    "SMH":   "Semiconductors",
    "XME":   "Metals & Mining",
    "ITA":   "Aerospace & Defense",
    "IBB":   "Biotech",
    "ARKK":  "Innovation / Growth",
    "ILF":   "Latin America",
    "EFV":   "Intl Value (EAFE)",
    "IWO":   "Small Cap Growth",
    "SDIV":  "Global SuperDividend",
    "URNM":  "Uranium / Nuclear",
    "COPX":  "Copper Miners",
    "PAVE":  "US Infrastructure",
    "CIBR":  "Cybersecurity",
    "AMLP":  "Midstream MLPs",
    "EMXC":  "EM ex-China",
    "GLDM":  "Physical Gold",
    "DBA":   "Broad Agriculture",
}

DROP_TICKERS = ["IGV", "IAI", "SOXX"]

ADD_TICKERS = {
    "URNM":  "Uranium / Nuclear",
    "COPX":  "Copper Miners",
    "PAVE":  "US Infrastructure",
    "CIBR":  "Cybersecurity",
    "AMLP":  "Midstream MLPs",
    "EMXC":  "EM ex-China",
    "GLDM":  "Physical Gold",
    "DBA":   "Broad Agriculture",
}

TLH_PAIRS = {
    "SMH":  "SOXX",
    "IBB":  "XBI",
    "SIL":  "SILJ",
    "XME":  "PICK",
    "ITA":  "XAR",
    "ARKK": "QQQJ",
    "IWO":  "VBK",
    "ILF":  "EWZ",
    "EFV":  "FNDF",
    "COPX": "CPER",
    "URNM": "URA",
    "PAVE": "IFRA",
    "CIBR": "HACK",
    "AMLP": "MLPA",
    "EMXC": "SCHE",
    "GLDM": "IAU",
    "DBA":  "PDBC",
}


# ═══════════════════════════════════════════════════════════════════════════════
# DATA & CORRELATION
# ═══════════════════════════════════════════════════════════════════════════════

def fetch_prices(tickers, start, end):
    unique = list(set(tickers))
    frames = []
    batch_size = 10
    for i in range(0, len(unique), batch_size):
        batch = unique[i:i + batch_size]
        data = yf.download(batch, start=start, end=end,
                           auto_adjust=True, progress=False)
        if data.empty:
            continue
        if isinstance(data.columns, pd.MultiIndex):
            frames.append(data["Close"])
        else:
            frames.append(data[["Close"]].rename(columns={"Close": batch[0]}))
    if not frames:
        return pd.DataFrame()
    prices = pd.concat(frames, axis=1)
    prices = prices.loc[:, ~prices.columns.duplicated()]
    return prices.ffill().dropna(how="all")


def corr_stats(returns, tickers):
    avail = [t for t in tickers if t in returns.columns]
    corr = returns[avail].corr()
    vals = corr.values[np.triu_indices_from(corr.values, k=1)]
    avg = np.nanmean(vals)
    median = np.nanmedian(vals)
    high_pairs = []
    cols = corr.columns.tolist()
    for i in range(len(cols)):
        for j in range(i + 1, len(cols)):
            c = corr.iloc[i, j]
            if abs(c) >= 0.70:
                high_pairs.append((cols[i], cols[j], round(c, 3)))
    high_pairs.sort(key=lambda x: -abs(x[2]))
    return avg, median, high_pairs, corr, avail


def candidate_corr_detail(returns, universe_tickers, candidate):
    if candidate not in returns.columns:
        return {}
    corrs = {}
    for t in universe_tickers:
        if t in returns.columns and t != candidate:
            corrs[t] = returns[[candidate, t]].corr().iloc[0, 1]
    return corrs


# ═══════════════════════════════════════════════════════════════════════════════
# WORD DOC HELPERS
# ═══════════════════════════════════════════════════════════════════════════════

def set_cell_shading(cell, hex_color):
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{hex_color}"/>')
    cell._tc.get_or_add_tcPr().append(shading)

def set_cell_border(cell, color="CCCCCC", size="4"):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    borders = parse_xml(
        f'<w:tcBorders {nsdecls("w")}>'
        f'  <w:top w:val="single" w:sz="{size}" w:space="0" w:color="{color}"/>'
        f'  <w:bottom w:val="single" w:sz="{size}" w:space="0" w:color="{color}"/>'
        f'  <w:left w:val="single" w:sz="{size}" w:space="0" w:color="{color}"/>'
        f'  <w:right w:val="single" w:sz="{size}" w:space="0" w:color="{color}"/>'
        f'</w:tcBorders>'
    )
    tcPr.append(borders)

def add_table(doc, headers, rows, col_widths=None, highlight_rows=None,
              highlight_color="E8F4E8", red_rows=None, red_color="FDEDED"):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False

    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        run.bold = True
        run.font.size = Pt(7.5)
        run.font.color.rgb = WHITE
        set_cell_shading(cell, "1A3C5E")
        set_cell_border(cell, color="1A3C5E")
        if col_widths:
            cell.width = Inches(col_widths[i])

    for r_idx, row_data in enumerate(rows):
        for c_idx, val in enumerate(row_data):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = ""
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if c_idx > 0 else WD_ALIGN_PARAGRAPH.LEFT
            run = p.add_run(str(val))
            run.font.size = Pt(7.5)
            if c_idx == 0:
                run.bold = True
            if highlight_rows and r_idx in highlight_rows:
                set_cell_shading(cell, highlight_color)
            if red_rows and r_idx in red_rows:
                set_cell_shading(cell, red_color)
            set_cell_border(cell, color="CCCCCC")
            if col_widths:
                cell.width = Inches(col_widths[c_idx])

    for row in table.rows:
        row.height = Cm(0.45)
        for cell in row.cells:
            for p in cell.paragraphs:
                p.paragraph_format.space_before = Pt(1)
                p.paragraph_format.space_after = Pt(1)
    return table

def add_navy_line(doc):
    line = doc.add_paragraph()
    line.paragraph_format.space_before = Pt(0)
    line.paragraph_format.space_after = Pt(2)
    pPr = line._p.get_or_add_pPr()
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'<w:bottom w:val="single" w:sz="6" w:space="1" w:color="1A3C5E"/>'
        f'</w:pBdr>'
    )
    pPr.append(pBdr)

def add_section_header(doc, text):
    h = doc.add_paragraph()
    run = h.add_run(text)
    run.bold = True
    run.font.size = Pt(9.5)
    run.font.color.rgb = NAVY
    h.paragraph_format.space_before = Pt(6)
    h.paragraph_format.space_after = Pt(2)

def add_body(doc, text, bold_prefix=None, size=9):
    p = doc.add_paragraph()
    if bold_prefix:
        run = p.add_run(bold_prefix)
        run.bold = True
        run.font.size = Pt(size)
    run = p.add_run(text)
    run.font.size = Pt(size)
    p.paragraph_format.space_after = Pt(2)
    return p

def add_footnote(doc, text):
    fn = doc.add_paragraph()
    run = fn.add_run(text)
    run.font.size = Pt(7)
    run.font.color.rgb = GRAY
    fn.paragraph_format.space_before = Pt(1)
    fn.paragraph_format.space_after = Pt(4)


# ═══════════════════════════════════════════════════════════════════════════════
# MAIN
# ═══════════════════════════════════════════════════════════════════════════════

def main():
    current_tickers = list(CURRENT_UNIVERSE.keys())
    proposed_tickers = list(PROPOSED_UNIVERSE.keys())
    all_tickers = list(set(current_tickers + proposed_tickers + ["SPY"]))

    print("Fetching price data...")
    prices = fetch_prices(all_tickers, "2022-01-01",
                          pd.Timestamp.today().strftime("%Y-%m-%d"))
    returns = prices.pct_change().dropna()
    if len(returns) > 504:
        returns = returns.iloc[-504:]

    print("Computing correlations...")
    c_avg, c_med, c_high, c_corr, c_avail = corr_stats(returns, current_tickers)
    p_avg, p_med, p_high, p_corr, p_avail = corr_stats(returns, proposed_tickers)

    spy_corrs_current = {t: returns[[t, "SPY"]].corr().iloc[0, 1]
                         for t in c_avail if "SPY" in returns.columns}
    spy_corrs_proposed = {t: returns[[t, "SPY"]].corr().iloc[0, 1]
                          for t in p_avail if "SPY" in returns.columns}

    print(f"  Current universe ({len(c_avail)}):  avg corr = {c_avg:.3f}")
    print(f"  Proposed universe ({len(p_avail)}): avg corr = {p_avg:.3f}")
    print(f"  Reduction: {p_avg - c_avg:+.3f}")

    # ── Build Word doc ───────────────────────────────────────────────────────
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10)
    style.paragraph_format.space_after = Pt(4)
    style.paragraph_format.space_before = Pt(0)
    for section in doc.sections:
        section.top_margin = Cm(1.3)
        section.bottom_margin = Cm(1.0)
        section.left_margin = Cm(1.8)
        section.right_margin = Cm(1.8)

    # Header
    h = doc.add_paragraph()
    h.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = h.add_run("POTOMAC FUND MANAGEMENT")
    run.bold = True
    run.font.size = Pt(15)
    run.font.color.rgb = NAVY
    h.paragraph_format.space_after = Pt(0)

    sub = doc.add_paragraph()
    run = sub.add_run("CRTOX Universe Expansion: Reducing Pairwise Correlation")
    run.bold = True
    run.font.size = Pt(12)
    run.font.color.rgb = DARK
    sub.paragraph_format.space_after = Pt(0)

    dateline = doc.add_paragraph()
    run = dateline.add_run("March 2026  |  Woody Wiegmann")
    run.font.size = Pt(10)
    run.font.color.rgb = GRAY
    dateline.paragraph_format.space_after = Pt(0)

    add_navy_line(doc)

    # ── Summary box ──────────────────────────────────────────────────────────
    add_section_header(doc, "Proposal")
    add_body(doc,
        f"Expand the CRTOX equity rotation universe from {len(current_tickers)} to "
        f"{len(proposed_tickers)} ETFs by adding 8 thematically distinct funds and "
        f"dropping 3 that overlap with existing holdings. "
        f"Average pairwise correlation falls from {c_avg:.3f} to {p_avg:.3f} "
        f"({p_avg - c_avg:+.3f}), broadening the opportunity set without changing "
        f"signals, rebalance frequency, or risk-off mechanics.")

    # ── Changes table ────────────────────────────────────────────────────────
    add_section_header(doc, "Universe Changes")

    add_headers = ["Add", "Ticker", "Theme", "Avg Corr\nto Universe", "SPY Corr", "Rationale"]
    add_rows = []
    for ticker, desc in ADD_TICKERS.items():
        cdetail = candidate_corr_detail(returns, [t for t in proposed_tickers if t != ticker], ticker)
        avg_c = np.mean(list(cdetail.values())) if cdetail else float("nan")
        spy_c = spy_corrs_proposed.get(ticker, float("nan"))
        rationales = {
            "URNM": "Uncorrelated nuclear/uranium theme",
            "COPX": "Pure copper exposure vs broad XME",
            "PAVE": "US infrastructure spending cycle",
            "CIBR": "Cybersecurity; distinct from broad tech",
            "AMLP": "Midstream MLPs; low equity beta, high yield",
            "EMXC": "EM without China concentration risk",
            "GLDM": "Physical gold; near-zero equity correlation",
            "DBA":  "Agriculture; weather-driven, non-equity factor",
        }
        add_rows.append([
            "+", ticker, desc,
            f"{avg_c:+.2f}" if not np.isnan(avg_c) else "N/A",
            f"{spy_c:+.2f}" if not np.isnan(spy_c) else "N/A",
            rationales.get(ticker, ""),
        ])

    add_table(doc, add_headers, add_rows,
              col_widths=[0.3, 0.5, 1.2, 0.7, 0.6, 2.9],
              highlight_rows=list(range(len(add_rows))), highlight_color="E8F4E8")

    doc.add_paragraph().paragraph_format.space_after = Pt(2)

    drop_headers = ["Drop", "Ticker", "Theme", "Reason"]
    drop_rows = []
    drop_reasons = {
        "IGV":  "Software; 0.85+ corr to QQQ, only held once in 6 months",
        "IAI":  "Broker-dealers; high overlap with financials/QQQ, only held once",
        "SOXX": "Semiconductors; 0.95+ corr to SMH (already in universe), never held",
    }
    for ticker in DROP_TICKERS:
        drop_rows.append([
            "-", ticker, CURRENT_UNIVERSE.get(ticker, ""),
            drop_reasons.get(ticker, ""),
        ])

    add_table(doc, drop_headers, drop_rows,
              col_widths=[0.3, 0.5, 1.2, 4.2],
              red_rows=list(range(len(drop_rows))), red_color="FDEDED")

    # ── Correlation comparison ───────────────────────────────────────────────
    add_section_header(doc, "Correlation Impact")

    corr_headers = ["Metric", "Current\n(13 ETFs)", "Proposed\n(18 ETFs)"]
    corr_rows = [
        ["Avg Pairwise Correlation", f"{c_avg:.3f}", f"{p_avg:.3f}"],
        ["Median Pairwise Correlation", f"{c_med:.3f}", f"{p_med:.3f}"],
        ["Pairs with r > 0.70", f"{len(c_high)}", f"{len(p_high)}"],
        ["Pairs with r > 0.85", f"{len([p for p in c_high if p[2] > 0.85])}",
         f"{len([p for p in p_high if p[2] > 0.85])}"],
    ]
    add_table(doc, corr_headers, corr_rows, col_widths=[2.5, 1.5, 1.5])

    # ── High corr pairs (current) ───────────────────────────────────────────
    if c_high:
        add_section_header(doc, f"Current Universe: High-Correlation Pairs (r > 0.70)")
        pair_headers = ["Pair", "Correlation", "Issue"]
        pair_rows = []
        for a, b, c in c_high[:10]:
            issue = ""
            if c > 0.85:
                issue = "Near-identical exposure"
            elif c > 0.75:
                issue = "Significant overlap"
            pair_rows.append([f"{a} / {b}", f"{c:+.3f}", issue])
        add_table(doc, pair_headers, pair_rows, col_widths=[1.5, 1.0, 3.7])

    # ── Proposed universe full listing ───────────────────────────────────────
    add_section_header(doc, "Proposed Universe (18 ETFs)")
    univ_headers = ["Ticker", "Theme", "SPY Corr", "Status"]
    univ_rows = []
    for ticker, desc in PROPOSED_UNIVERSE.items():
        spy_c = spy_corrs_proposed.get(ticker, float("nan"))
        if ticker in ADD_TICKERS:
            status = "NEW"
        else:
            status = "Keep"
        univ_rows.append([
            ticker, desc,
            f"{spy_c:+.2f}" if not np.isnan(spy_c) else "N/A",
            status,
        ])
    new_rows = [i for i, r in enumerate(univ_rows) if r[3] == "NEW"]
    add_table(doc, univ_headers, univ_rows, col_widths=[0.6, 1.6, 0.7, 0.6],
              highlight_rows=new_rows, highlight_color="E8F4E8")

    # ── TLH pairs ────────────────────────────────────────────────────────────
    add_section_header(doc, "Tax-Loss Harvesting Pairs")
    add_body(doc,
        "TLH-only instruments (not in the rotation universe). On each momentum-driven "
        "rotation, screen exiting lots for unrealized losses and swap into the paired ETF "
        "to maintain exposure while harvesting the loss.",
        size=8)

    tlh_headers = ["Primary", "TLH Swap", "Exposure"]
    tlh_rows = []
    exposure_map = {
        "SMH": "Semiconductors", "IBB": "Biotech", "SIL": "Silver miners",
        "XME": "Metals & mining", "ITA": "Aerospace & defense",
        "ARKK": "Innovation / growth", "IWO": "Small cap growth",
        "ILF": "Latin America", "EFV": "Intl value",
        "COPX": "Copper miners", "URNM": "Uranium",
        "PAVE": "US infrastructure", "CIBR": "Cybersecurity",
        "AMLP": "Midstream MLPs", "EMXC": "EM ex-China",
        "GLDM": "Physical gold", "DBA": "Agriculture",
    }
    for primary, swap in TLH_PAIRS.items():
        tlh_rows.append([primary, swap, exposure_map.get(primary, "")])
    add_table(doc, tlh_headers, tlh_rows, col_widths=[0.8, 0.8, 2.0])

    # ── Footer ───────────────────────────────────────────────────────────────
    add_navy_line(doc)

    bl = doc.add_paragraph()
    run = bl.add_run("Bottom Line: ")
    run.bold = True
    run.font.size = Pt(9)
    run.font.color.rgb = NAVY
    run = bl.add_run(
        "Same signals, same rebalance cadence, same risk-off mechanics. "
        "The only change is what the fund can choose from. A broader, less correlated "
        "menu gives the momentum algorithm more room to find genuinely differentiated opportunities."
    )
    run.font.size = Pt(9)
    run.font.color.rgb = DARK

    end = returns.index[-1].strftime("%Y-%m-%d")
    start = returns.index[0].strftime("%Y-%m-%d")
    add_footnote(doc,
        f"Correlations computed on daily returns, {start} to {end} (~2 years). "
        f"Source: Yahoo Finance. Past correlations do not guarantee future correlations.")

    # Save
    os.makedirs(OUT_DIR, exist_ok=True)
    out_path = os.path.join(OUT_DIR, "CRTOX Universe Expansion - Mar 2026.docx")
    doc.save(out_path)
    print(f"\nSaved to: {out_path}")

    out_path2 = os.path.join(SCRIPT_DIR, "CRTOX Universe Expansion - Mar 2026.docx")
    doc.save(out_path2)
    print(f"Also saved to: {out_path2}")


if __name__ == "__main__":
    main()
