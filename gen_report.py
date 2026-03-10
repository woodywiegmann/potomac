"""Generate the CRTPX Backtest Report Word document from results + trade CSV."""

import json
import csv
import os
from datetime import datetime
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml

SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))

NAVY = RGBColor(0x1A, 0x3C, 0x5E)
DARK = RGBColor(0x33, 0x33, 0x33)
GRAY = RGBColor(0x66, 0x66, 0x66)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)

QC_URL = "https://www.quantconnect.com/terminal/28552769#open/f4059b5808096b3fc4a98ae8da4bee11"


def shade(cell, hex_color):
    cell._tc.get_or_add_tcPr().append(
        parse_xml(f'<w:shd {nsdecls("w")} w:fill="{hex_color}"/>')
    )


def border_cell(cell, color="CCCCCC"):
    cell._tc.get_or_add_tcPr().append(parse_xml(
        f'<w:tcBorders {nsdecls("w")}>'
        f'<w:top w:val="single" w:sz="4" w:space="0" w:color="{color}"/>'
        f'<w:bottom w:val="single" w:sz="4" w:space="0" w:color="{color}"/>'
        f'<w:left w:val="single" w:sz="4" w:space="0" w:color="{color}"/>'
        f'<w:right w:val="single" w:sz="4" w:space="0" w:color="{color}"/>'
        f'</w:tcBorders>'
    ))


def add_table(doc, headers, rows, col_widths=None, highlight_col=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        run.bold = True
        run.font.size = Pt(8)
        run.font.color.rgb = WHITE
        shade(cell, "1A3C5E")
        border_cell(cell, "1A3C5E")
        if col_widths:
            cell.width = Inches(col_widths[i])
    for r_idx, row_data in enumerate(rows):
        for c_idx, val in enumerate(row_data):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = ""
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if c_idx > 0 else WD_ALIGN_PARAGRAPH.LEFT
            run = p.add_run(str(val))
            run.font.size = Pt(8)
            if c_idx == 0:
                run.bold = True
            if highlight_col is not None and c_idx == highlight_col:
                shade(cell, "E8F4E8")
            border_cell(cell, "CCCCCC")
            if col_widths:
                cell.width = Inches(col_widths[c_idx])
    for row in table.rows:
        row.height = Cm(0.5)
        for cell in row.cells:
            for p in cell.paragraphs:
                p.paragraph_format.space_before = Pt(1)
                p.paragraph_format.space_after = Pt(1)
    return table


def analyze_trades(csv_path):
    equity_tickers = {"VOO", "IVV", "SPLG"}
    cash_tickers = {"SGOV", "BIL", "SHV"}

    trades = []
    with open(csv_path) as f:
        for row in csv.DictReader(f):
            trades.append(row)

    regime = None
    switches = 0
    tlh_swaps = 0
    on_periods = []
    off_periods = []
    current_start = None

    for t in trades:
        if t["action"] != "BUY":
            continue
        ticker = t["ticker"]
        date = t["date"]

        if ticker in equity_tickers:
            new_regime = "ON"
        elif ticker in cash_tickers:
            new_regime = "OFF"
        else:
            continue

        if regime is None:
            regime = new_regime
            current_start = date
        elif new_regime != regime:
            switches += 1
            if current_start:
                if regime == "ON":
                    on_periods.append((current_start, date))
                else:
                    off_periods.append((current_start, date))
            regime = new_regime
            current_start = date
        else:
            tlh_swaps += 1

    total_on = sum(
        (datetime.strptime(e, "%Y-%m-%d") - datetime.strptime(s, "%Y-%m-%d")).days
        for s, e in on_periods
    )
    total_off = sum(
        (datetime.strptime(e, "%Y-%m-%d") - datetime.strptime(s, "%Y-%m-%d")).days
        for s, e in off_periods
    )
    total_days = total_on + total_off
    years = total_days / 365.25 if total_days > 0 else 1

    return {
        "total_trades": len(trades),
        "switches": switches,
        "tlh_swaps": tlh_swaps,
        "switches_per_year": round(switches / years, 1),
        "on_days": total_on,
        "off_days": total_off,
        "on_pct": round(total_on / total_days * 100) if total_days else 0,
        "off_pct": round(total_off / total_days * 100) if total_days else 0,
    }


def main():
    results_path = os.path.join(SCRIPT_DIR, "qc_crtpx_results.json")
    csv_path = os.path.join(SCRIPT_DIR, "crtpx_trades.csv")

    with open(results_path) as f:
        data = json.load(f)
    bt = data.get("backtest", {})
    stats = bt.get("statistics", {})
    runtime = bt.get("runtimeStatistics", {})

    trade_stats = analyze_trades(csv_path)

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10)
    style.paragraph_format.space_after = Pt(4)
    for section in doc.sections:
        section.top_margin = Cm(1.5)
        section.bottom_margin = Cm(1.2)
        section.left_margin = Cm(2.0)
        section.right_margin = Cm(2.0)

    # ── HEADER ──
    h = doc.add_paragraph()
    run = h.add_run("POTOMAC FUND MANAGEMENT")
    run.bold = True
    run.font.size = Pt(16)
    run.font.color.rgb = NAVY
    h.paragraph_format.space_after = Pt(0)

    sub = doc.add_paragraph()
    run = sub.add_run("CRTPX Strategy B: 1.2x ES Futures + Conditional Risk-Off")
    run.bold = True
    run.font.size = Pt(13)
    run.font.color.rgb = DARK
    sub.paragraph_format.space_after = Pt(0)

    dateline = doc.add_paragraph()
    run = dateline.add_run(
        "February 2026  |  QuantConnect Backtest  |  June 2019 \u2013 February 2026"
    )
    run.font.size = Pt(10)
    run.font.color.rgb = GRAY
    dateline.paragraph_format.space_after = Pt(0)

    author = doc.add_paragraph()
    run = author.add_run("Woody Wiegmann")
    run.font.size = Pt(10)
    run.font.color.rgb = GRAY
    author.paragraph_format.space_after = Pt(2)

    line = doc.add_paragraph()
    line.paragraph_format.space_before = Pt(0)
    line.paragraph_format.space_after = Pt(4)
    pPr = line._p.get_or_add_pPr()
    pPr.append(parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'<w:bottom w:val="single" w:sz="6" w:space="1" w:color="1A3C5E"/>'
        f'</w:pBdr>'
    ))

    # ── SIGNAL ARCHITECTURE ──
    sig_h = doc.add_paragraph()
    run = sig_h.add_run("Signal Architecture")
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = NAVY
    sig_h.paragraph_format.space_after = Pt(2)

    signals = [
        ("Penta1", "SPY close > 50-day SMA(SPY)", "Trend"),
        ("Penta2", "IYT close > 50-day SMA(IYT)", "Economic confirmation"),
        ("Penta3", "^NYA close > 50-day SMA(^NYA)", "NYSE breadth (index data)"),
        ("Penta4", "LQD close > 50-day SMA(LQD)", "Credit conditions"),
    ]
    for name, desc, label in signals:
        p = doc.add_paragraph()
        run = p.add_run(f"{name}: ")
        run.bold = True
        run.font.size = Pt(9)
        run = p.add_run(f"{desc}  ({label})")
        run.font.size = Pt(9)
        p.paragraph_format.space_after = Pt(1)
        p.paragraph_format.left_indent = Inches(0.15)

    conf_p = doc.add_paragraph()
    run = conf_p.add_run(
        "Penta ON = 3+ of 4 green.  3-day confirmation (Sum/ExRem).  "
        "Risk-on: 1.2x S&P 500 via ES futures (excess cash earns T-bill yield)."
    )
    run.font.size = Pt(9)
    run.bold = True
    conf_p.paragraph_format.space_before = Pt(4)
    conf_p.paragraph_format.space_after = Pt(2)

    roff_h = doc.add_paragraph()
    run = roff_h.add_run("Conditional Risk-Off Allocation")
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = NAVY
    roff_h.paragraph_format.space_after = Pt(2)

    roff_p1 = doc.add_paragraph()
    run = roff_p1.add_run("Bearish (SPY < 50-day SMA): ")
    run.bold = True
    run.font.size = Pt(9)
    run = roff_p1.add_run("50% SGOV + 50% CAOS")
    run.font.size = Pt(9)
    roff_p1.paragraph_format.space_after = Pt(1)
    roff_p1.paragraph_format.left_indent = Inches(0.15)

    roff_p2 = doc.add_paragraph()
    run = roff_p2.add_run("Bullish (SPY > 50-day SMA): ")
    run.bold = True
    run.font.size = Pt(9)
    run = roff_p2.add_run("50% SGOV + 25% CAOS + 25% DBMF")
    run.font.size = Pt(9)
    roff_p2.paragraph_format.space_after = Pt(2)
    roff_p2.paragraph_format.left_indent = Inches(0.15)

    roff_rationale = doc.add_paragraph()
    run = roff_rationale.add_run(
        "Risk-on leverage: ES futures provide 1.2x notional S&P 500 exposure with margin "
        "efficiency. Excess cash earns T-bill yield implicitly. Auto-rolls to front month contract. "
        "Risk-off (1.0x, no leverage): CAOS provides convex put-like payoff "
        "at drawdown onset. DBMF replicates hedge fund managed futures (bullish risk-off only). "
        "SGOV anchors 50% of capital with zero equity beta."
    )
    run.font.size = Pt(8)
    run.font.color.rgb = GRAY
    roff_rationale.paragraph_format.space_after = Pt(6)

    # ── BACKTEST RESULTS ──
    sum_h = doc.add_paragraph()
    run = sum_h.add_run("Backtest Results")
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = NAVY
    sum_h.paragraph_format.space_after = Pt(2)

    sum_rows = [
        ["Period", "June 2019 \u2013 February 2026 (6.7 years)"],
        ["Starting Capital", "$1,000,000"],
        ["Final Equity", runtime.get("Equity", "N/A")],
        ["CAGR", stats.get("Compounding Annual Return", "N/A")],
        ["Net Profit", stats.get("Net Profit", "N/A")],
        ["Max Drawdown", stats.get("Drawdown", "N/A")],
        ["Sharpe Ratio", stats.get("Sharpe Ratio", "N/A")],
        ["Sortino Ratio", stats.get("Sortino Ratio", "N/A")],
        ["Alpha", stats.get("Alpha", "N/A")],
        ["Beta", stats.get("Beta", "N/A")],
        ["Win Rate", stats.get("Win Rate", "N/A")],
        ["Loss Rate", stats.get("Loss Rate", "N/A")],
        ["Avg Win", stats.get("Average Win", "N/A")],
        ["Avg Loss", stats.get("Average Loss", "N/A")],
        ["Total Orders", stats.get("Total Orders", "N/A")],
        ["Total Fees", stats.get("Total Fees", "N/A")],
        ["Portfolio Turnover", stats.get("Portfolio Turnover", "N/A")],
    ]
    add_table(doc, ["Metric", "Value"], sum_rows, col_widths=[2.2, 3.0])

    # ── TRADE ACTIVITY ──
    act_h = doc.add_paragraph()
    run = act_h.add_run("Trade Activity (from Trade CSV)")
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = NAVY
    act_h.paragraph_format.space_before = Pt(8)
    act_h.paragraph_format.space_after = Pt(2)

    ts = trade_stats
    act_rows = [
        ["Regime switches", str(ts["switches"])],
        ["Switches/year", str(ts["switches_per_year"])],
        ["TLH swaps", str(ts["tlh_swaps"])],
        ["Risk-on days", f'{ts["on_days"]:,} ({ts["on_pct"]}%)'],
        ["Risk-off days", f'{ts["off_days"]:,} ({ts["off_pct"]}%)'],
        ["Total trades", str(ts["total_trades"])],
    ]
    add_table(doc, ["Metric", "Value"], act_rows, col_widths=[2.2, 3.0])

    # ── AMIBROKER CODE ──
    ami_h = doc.add_paragraph()
    run = ami_h.add_run("AmiBroker Implementation (for Developer)")
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = NAVY
    ami_h.paragraph_format.space_before = Pt(10)
    ami_h.paragraph_format.space_after = Pt(4)

    ami_code = [
        'SMA_Period = 50;  ConfirmDays = 3;',
        '',
        '// ---- Penta signal indicators ----',
        'Penta1 = Foreign("SPY","C") > MA(Foreign("SPY","C"), SMA_Period);',
        'Penta2 = Foreign("IYT","C") > MA(Foreign("IYT","C"), SMA_Period);',
        'Penta3 = Foreign("~NYA","C") > MA(Foreign("~NYA","C"), SMA_Period);',
        'Penta4 = Foreign("LQD","C") > MA(Foreign("LQD","C"), SMA_Period);',
        '',
        'PentaScore = Penta1 + Penta2 + Penta3 + Penta4;',
        'RawOn = PentaScore >= 3;',
        '',
        '// 3-day confirmation',
        'ConfOn  = Sum(RawOn, ConfirmDays) == ConfirmDays;',
        'ConfOff = Sum(!RawOn, ConfirmDays) == ConfirmDays;',
        'Buy  = ExRem(ConfOn, ConfOff);',
        'Sell = ExRem(ConfOff, ConfOn);',
        '',
        '// ---- Risk-on allocation ----',
        '// 1.2x S&P 500 notional via ES futures',
        '// (AmiBroker: use SetPositionSize(120, spsPercentOfEquity)',
        '//  or track ES continuous front-month contract)',
        '',
        '// ---- Conditional risk-off (Strategy B) ----',
        'SpyBearish = Foreign("SPY","C") < MA(Foreign("SPY","C"), SMA_Period);',
        '',
        '// Bearish risk-off (SPY below 50-day SMA):',
        'RiskOffBearRet = 0.50 * ROC(Foreign("SGOV","C"),1)/100',
        '               + 0.50 * ROC(Foreign("CAOS","C"),1)/100;',
        '',
        '// Bullish risk-off (SPY above 50-day SMA):',
        'RiskOffBullRet = 0.50 * ROC(Foreign("SGOV","C"),1)/100',
        '               + 0.25 * ROC(Foreign("CAOS","C"),1)/100',
        '               + 0.25 * ROC(Foreign("DBMF","C"),1)/100;',
        '',
        '// Select blend based on SPY vs SMA',
        'RiskOffRet = IIf(SpyBearish, RiskOffBearRet, RiskOffBullRet);',
        '',
        '// No TLH needed -- futures are Section 1256 contracts',
        '// (60/40 long-term/short-term tax treatment)',
    ]
    for codeline in ami_code:
        p = doc.add_paragraph()
        run = p.add_run(codeline)
        run.font.size = Pt(8)
        run.font.name = "Consolas"
        run.font.color.rgb = DARK
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.left_indent = Inches(0.3)

    # ── SENSITIVITY ──
    sens_h = doc.add_paragraph()
    run = sens_h.add_run("Sensitivity Testing Parameters")
    run.bold = True
    run.font.size = Pt(10)
    run.font.color.rgb = NAVY
    sens_h.paragraph_format.space_before = Pt(8)
    sens_h.paragraph_format.space_after = Pt(2)

    for item in [
        "SMA period: test 40, 50, 60, 80",
        "Confirmation days: test 2, 3, 5",
        "TLH loss threshold: test -2%, -3%, -5%",
        "CTA vs DBMF bearish/bullish split: test alternative thresholds",
        "CAOS weight: test 20-40% range",
    ]:
        p = doc.add_paragraph()
        run = p.add_run(f"\u2022 {item}")
        run.font.size = Pt(9)
        p.paragraph_format.space_after = Pt(1)
        p.paragraph_format.left_indent = Inches(0.15)

    overfit_p = doc.add_paragraph()
    run = overfit_p.add_run("Why this isn't overfit: ")
    run.bold = True
    run.font.size = Pt(9)
    run.font.color.rgb = NAVY
    run = overfit_p.add_run(
        "4 standard indicators, 1 parameter each (50-day SMA). "
        "3-day confirmation is a standard anti-whipsaw filter. "
        "Only 2 free parameters total."
    )
    run.font.size = Pt(9)
    overfit_p.paragraph_format.space_before = Pt(4)

    # ── FOOTER ──
    footer_line = doc.add_paragraph()
    footer_line.paragraph_format.space_before = Pt(10)
    footer_line.paragraph_format.space_after = Pt(2)
    pPr = footer_line._p.get_or_add_pPr()
    pPr.append(parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'<w:bottom w:val="single" w:sz="6" w:space="1" w:color="1A3C5E"/>'
        f'</w:pBdr>'
    ))

    bl = doc.add_paragraph()
    run = bl.add_run("Data Sources: ")
    run.bold = True
    run.font.size = Pt(8)
    run.font.color.rgb = NAVY
    run = bl.add_run(
        "SPY (equity), IYT (equity), NYA (QuantConnect Cash Index \u2014 actual NYSE Composite), "
        "LQD (equity). SMA computed by QuantConnect built-in SMA indicator on daily close prices. "
        "Trade CSV (crtpx_trades.csv) contains all 107 orders with dates, prices, and share counts."
    )
    run.font.size = Pt(8)
    run.font.color.rgb = DARK

    link_p = doc.add_paragraph()
    run = link_p.add_run(f"QuantConnect backtest: {QC_URL}")
    run.font.size = Pt(8)
    run.font.color.rgb = GRAY

    out_path = os.path.join(SCRIPT_DIR, "CRTPX_StratB_Backtest_Report.docx")
    doc.save(out_path)
    print(f"Saved: {out_path}")


if __name__ == "__main__":
    main()
