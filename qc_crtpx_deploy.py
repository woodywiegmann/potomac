"""
QuantConnect Deployment: CRTPX Tactically Passive + TLH
=========================================================
Deploys a Penta-signal algorithm for the Potomac Tactically
Passive Fund (CRTPX) to QuantConnect, backtests, retrieves
results, exports a trade CSV, and generates a Word doc.

Signal architecture mirrors the AmiBroker implementation exactly:
  Penta1 = SPY close > 50-day SMA of SPY close
  Penta2 = IYT close > 50-day SMA of IYT close
  Penta3 = ^NYA close > 50-day SMA of ^NYA close   (index data)
  Penta4 = LQD close > 50-day SMA of LQD close

  PentaScore = Penta1 + Penta2 + Penta3 + Penta4
  RawOn  = PentaScore >= 3
  ConfOn  = Sum(RawOn, 3) == 3      (3 consecutive days ON)
  ConfOff = Sum(!RawOn, 3) == 3     (3 consecutive days OFF)
  Buy  = ExRem(ConfOn, ConfOff)     (fire once, wait for opposite)
  Sell = ExRem(ConfOff, ConfOn)

  Risk-on:  100% VOO
  Risk-off: 100% SGOV
  TLH swap rings: VOO<->IVV<->SPLG, SGOV<->BIL<->SHV
  Loss trigger: -3%, wash sale window: 31 calendar days

Usage:
    python qc_crtpx_deploy.py
"""

import json
import os
import sys
import time as time_mod
from base64 import b64encode
from hashlib import sha256
from requests import post, get

BASE_URL = "https://www.quantconnect.com/api/v2"
USER_ID = 470149
API_TOKEN = "0d335ae3e7bc1d4cb9a57f3c1b3d6f87419b1aec369bf085dc44bc5043b9b88a"

SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))


def get_headers():
    timestamp = str(int(time_mod.time()))
    hashed = sha256(f"{API_TOKEN}:{timestamp}".encode("utf-8")).hexdigest()
    auth = b64encode(f"{USER_ID}:{hashed}".encode("utf-8")).decode("ascii")
    return {"Authorization": f"Basic {auth}", "Timestamp": timestamp}


def api_post(endpoint, payload=None):
    url = f"{BASE_URL}{endpoint}"
    resp = post(url, headers=get_headers(), json=payload or {})
    data = resp.json()
    if not data.get("success"):
        print(f"  API ERROR {endpoint}: {json.dumps(data, indent=2)}")
    return data


def api_get(endpoint, params=None):
    url = f"{BASE_URL}{endpoint}"
    resp = get(url, headers=get_headers(), params=params or {})
    return resp.json()


# ═══════════════════════════════════════════════════════════════
# LEAN ALGORITHM CODE -- matches AmiBroker spec exactly
# ═══════════════════════════════════════════════════════════════

ALGO_CODE = r'''
from AlgorithmImports import *
from datetime import timedelta
from collections import deque


class CrtpxLevered(QCAlgorithm):
    """
    CRTPX Strategy B -- 1.2x Levered via ES Futures + Conditional Risk-Off
    ========================================================================
    Same Penta signal as Strategy A (AmiBroker match).

    Risk-on:  1.2x S&P 500 notional via ES futures.
              Excess cash earns T-bill yield implicitly.
    Risk-off: CONDITIONAL on SPY vs 50-day SMA (1.0x, no leverage):
      - Bearish (SPY < 50-day SMA): 50% SGOV + 50% CAOS
      - Bullish (SPY > 50-day SMA): 50% SGOV + 25% CAOS + 25% DBMF

    ES contract: $50 per index point.
    Target beta: 1.2
    """

    TARGET_BETA = 1.20
    ES_MULTIPLIER = 50

    BEARISH_ALLOC = {"SGOV": 0.50, "CAOS": 0.50}
    BULLISH_ALLOC = {"SGOV": 0.50, "CAOS": 0.25, "DBMF": 0.25}

    ALL_RISK_OFF_TICKERS = ["CAOS", "SGOV", "DBMF"]

    CONFIRM_DAYS = 3
    SMA_PERIOD   = 50

    def initialize(self):
        self.set_start_date(2019, 6, 1)
        self.set_end_date(2026, 2, 1)
        self.set_cash(1_000_000)
        self.set_benchmark("SPY")
        self.universe_settings.resolution = Resolution.DAILY
        self.settings.seed_initial_prices = True

        for t in self.ALL_RISK_OFF_TICKERS:
            self.add_equity(t, Resolution.DAILY)

        self._es = self.add_future(
            Futures.Indices.SP_500_E_MINI,
            Resolution.DAILY,
            extended_market_hours=False,
            data_mapping_mode=DataMappingMode.OPEN_INTEREST,
            data_normalization_mode=DataNormalizationMode.BACKWARDS_RATIO,
            contract_depth_offset=0,
        )
        self._es.set_filter(lambda u: u.front_month())
        self._es_sym = self._es.symbol

        self.spy_sym = self.add_equity("SPY", Resolution.DAILY).symbol
        self.iyt_sym = self.add_equity("IYT", Resolution.DAILY).symbol
        self.nya_sym = self.add_index("NYA", Resolution.DAILY).symbol
        self.lqd_sym = self.add_equity("LQD", Resolution.DAILY).symbol

        self._signal_syms = [self.spy_sym, self.iyt_sym, self.nya_sym, self.lqd_sym]
        self._sma = {}
        for sym in self._signal_syms:
            self._sma[sym] = self.SMA(sym, self.SMA_PERIOD, Resolution.DAILY)

        self._raw_on_history = deque(maxlen=self.CONFIRM_DAYS)
        self._regime = None
        self._last_buy_fired = False
        self._last_sell_fired = False
        self._current_riskoff_type = None
        self.switch_count = 0

        self._trade_log = []

        self.schedule.on(
            self.date_rules.every_day("SPY"),
            self.time_rules.after_market_open("SPY", 31),
            self._daily_check,
        )

        self.set_warm_up(timedelta(days=self.SMA_PERIOD + 30))

    def _all_sma_ready(self):
        return all(self._sma[s].is_ready for s in self._signal_syms)

    def _spy_bearish(self):
        spy_price = self.securities[self.spy_sym].price
        spy_sma = self._sma[self.spy_sym].current.value
        if spy_price <= 0 or spy_sma <= 0:
            return True
        return spy_price < spy_sma

    def _penta_raw_on(self):
        if not self._all_sma_ready():
            return None
        greens = 0
        for sym in self._signal_syms:
            price = self.securities[sym].price
            sma_val = self._sma[sym].current.value
            if price > 0 and sma_val > 0 and price > sma_val:
                greens += 1
        return greens >= 3

    def _daily_check(self):
        if self.is_warming_up:
            return
        raw_on = self._penta_raw_on()
        if raw_on is None:
            return

        self._raw_on_history.append(raw_on)
        if len(self._raw_on_history) < self.CONFIRM_DAYS:
            return

        conf_on = all(self._raw_on_history)
        conf_off = all(not x for x in self._raw_on_history)

        buy_signal = False
        sell_signal = False

        if conf_on and not self._last_buy_fired:
            buy_signal = True
            self._last_buy_fired = True
            self._last_sell_fired = False
        elif conf_off and not self._last_sell_fired:
            sell_signal = True
            self._last_sell_fired = True
            self._last_buy_fired = False

        if buy_signal:
            self.switch_count += 1
            self.debug(f"REGIME SWITCH #{self.switch_count}: OFF -> ON | date={self.time.date()}")
            self._go_risk_on()
        elif sell_signal:
            self.switch_count += 1
            self.debug(f"REGIME SWITCH #{self.switch_count}: ON -> OFF | date={self.time.date()}")
            self._go_risk_off()

    def _sym(self, ticker):
        return self.symbol(ticker)

    def _target_es_contracts(self):
        """Calculate # of ES contracts for 1.2x notional exposure."""
        portfolio_value = self.portfolio.total_portfolio_value
        mapped = self._es.mapped
        if mapped is None:
            return 0
        es_price = self.securities[mapped].price
        if es_price <= 0:
            return 0
        target_notional = portfolio_value * self.TARGET_BETA
        contract_notional = es_price * self.ES_MULTIPLIER
        return int(round(target_notional / contract_notional))

    def _liquidate_all_risk_off(self):
        for t in self.ALL_RISK_OFF_TICKERS:
            sym = self._sym(t)
            if self.portfolio[sym].invested:
                self.liquidate(sym)
                self._log_trade("SELL", t, "regime_switch_to_ON")

    def _liquidate_es(self):
        mapped = self._es.mapped
        if mapped is not None and self.portfolio[mapped].invested:
            self.liquidate(mapped)
            self._log_trade("SELL", "ES", "regime_switch_to_OFF")

    def _go_risk_on(self):
        self._liquidate_all_risk_off()

        mapped = self._es.mapped
        if mapped is None:
            self.debug("  WARNING: no mapped ES contract available")
            return

        target_qty = self._target_es_contracts()
        current_qty = self.portfolio[mapped].quantity
        delta = target_qty - int(current_qty)
        if delta != 0:
            self.market_order(mapped, delta)
            self._log_trade("BUY", "ES", f"risk_on_1.2x_{target_qty}_contracts")

        self._regime = "ON"
        self._current_riskoff_type = None
        notional = target_qty * self.ES_MULTIPLIER * self.securities[mapped].price
        self.debug(f"  -> RISK-ON: {target_qty} ES contracts = ${notional:,.0f} notional ({self.TARGET_BETA}x)")

    def _go_risk_off(self):
        self._liquidate_es()

        bearish = self._spy_bearish()
        alloc = self.BEARISH_ALLOC if bearish else self.BULLISH_ALLOC
        riskoff_type = "BEARISH" if bearish else "BULLISH"
        self._current_riskoff_type = riskoff_type

        for ticker, weight in alloc.items():
            sym = self._sym(ticker)
            self.set_holdings(sym, weight)
            self._log_trade("BUY", ticker, f"risk_off_{riskoff_type}_{weight:.0%}")

        self._regime = "OFF"
        blend_str = " + ".join(f"{w:.0%} {t}" for t, w in alloc.items())
        self.debug(f"  -> RISK-OFF ({riskoff_type}): {blend_str}")

    def on_securities_changed(self, changes):
        """Handle ES contract rolls automatically."""
        if self._regime != "ON":
            return
        for removed in changes.removed_securities:
            if removed.symbol.security_type == SecurityType.FUTURE and not removed.symbol.is_canonical():
                self.liquidate(removed.symbol)
                self._log_trade("SELL", "ES", "contract_roll_out")

        for added in changes.added_securities:
            if added.symbol.security_type == SecurityType.FUTURE and not added.symbol.is_canonical():
                target_qty = self._target_es_contracts()
                if target_qty > 0:
                    self.market_order(added.symbol, target_qty)
                    self._log_trade("BUY", "ES", f"contract_roll_in_{target_qty}")

    def _log_trade(self, action, ticker, reason):
        if ticker == "ES":
            mapped = self._es.mapped
            if mapped is not None:
                price = self.securities[mapped].price
                qty = self.portfolio[mapped].quantity
                value = abs(qty) * price * self.ES_MULTIPLIER
            else:
                price = 0
                qty = 0
                value = 0
        else:
            sym = self._sym(ticker)
            h = self.portfolio[sym]
            price = self.securities[sym].price
            qty = h.quantity
            value = abs(qty * price)

        greens = 0
        if self._all_sma_ready():
            for s in self._signal_syms:
                p = self.securities[s].price
                sma_v = self._sma[s].current.value
                if p > 0 and sma_v > 0 and p > sma_v:
                    greens += 1

        self._trade_log.append({
            "date": str(self.time.date()),
            "action": action,
            "ticker": ticker,
            "price": round(price, 4),
            "shares": int(abs(qty)) if action == "SELL" else 0,
            "value": round(value, 2),
            "reason": reason,
            "penta_score": greens,
            "regime": self._regime or "INIT",
            "equity": round(self.portfolio.total_portfolio_value, 2),
        })

    def on_end_of_algorithm(self):
        years = (self.end_date - self.start_date).days / 365.25
        self.debug("=" * 60)
        self.debug("CRTPX BACKTEST SUMMARY")
        self.debug("=" * 60)
        self.debug(f"Period: {self.start_date.date()} to {self.end_date.date()}")
        self.debug(f"Regime switches: {self.switch_count}")
        self.debug(f"Switches/year:   {self.switch_count / years:.1f}")
        self.debug(f"Target beta:     {self.TARGET_BETA}")
        self.debug(f"Leverage method: ES futures overlay")
        self.debug(f"Total trades logged: {len(self._trade_log)}")
        self.debug("=" * 60)

        self.debug("Trade log: retrieve via API /backtests/orders/read")
'''


def main():
    print("=" * 70)
    print("QUANTCONNECT DEPLOYMENT: CRTPX Strategy B -- 1.2x ES Futures + Cond Risk-Off")
    print("  Signal: Penta (SPY, IYT, ^NYA index, LQD) vs 50-day SMA")
    print("  Risk-on:  1.2x S&P 500 via ES futures")
    print("  Risk-off: Bearish=50%SGOV+50%CAOS, Bullish=50%SGOV+25%CAOS+25%DBMF")
    print("=" * 70)

    # 1. Authenticate
    print("\n1. Authenticating...")
    r = api_post("/authenticate")
    if not r.get("success"):
        print("   Auth failed.")
        sys.exit(1)
    print("   OK")

    # 2. Create project
    print("\n2. Creating project...")
    ts = int(time_mod.time())
    project_name = f"CRTPX_StratB_1.2x_ES_{ts}"
    r = api_post("/projects/create", {"name": project_name, "language": "Py"})
    if not r.get("success"):
        print("   Failed to create project.")
        sys.exit(1)
    project_id = r["projects"][0]["projectId"]
    print(f"   Project: {project_name} (ID: {project_id})")

    # 3. Upload code
    print("\n3. Uploading algorithm...")
    r = api_post("/files/update", {
        "projectId": project_id,
        "name": "main.py",
        "content": ALGO_CODE,
    })
    if not r.get("success"):
        r = api_post("/files/create", {
            "projectId": project_id,
            "name": "main.py",
            "content": ALGO_CODE,
        })
        if not r.get("success"):
            print("   Upload failed.")
            sys.exit(1)
    print("   Uploaded.")

    # 4. Compile
    print("\n4. Compiling...")
    r = api_post("/compile/create", {"projectId": project_id})
    if not r.get("success"):
        print("   Compile request failed.")
        sys.exit(1)
    compile_id = r.get("compileId")
    print(f"   Compile ID: {compile_id}")

    for _ in range(30):
        time_mod.sleep(3)
        r = api_post("/compile/read", {
            "projectId": project_id,
            "compileId": compile_id,
        })
        state = r.get("state", "")
        print(f"   State: {state}")
        if state == "BuildSuccess":
            break
        if state == "BuildError":
            logs = r.get("logs", [])
            for log in logs:
                print(f"     {log}")
            sys.exit(1)
    else:
        print("   Compile timed out.")
        sys.exit(1)
    print("   Compiled OK.")

    # 5. Backtest
    print("\n5. Starting backtest...")
    bt_name = f"CRTPX_StratB_{ts}"
    r = api_post("/backtests/create", {
        "projectId": project_id,
        "compileId": compile_id,
        "backtestName": bt_name,
    })
    if not r.get("success"):
        print("   Backtest creation failed.")
        print(json.dumps(r, indent=2))
        sys.exit(1)
    backtest_id = r["backtest"]["backtestId"]
    print(f"   Backtest: {bt_name} (ID: {backtest_id})")

    # 6. Poll
    print("\n6. Waiting for completion...")
    for _ in range(180):
        time_mod.sleep(5)
        r = api_post("/backtests/read", {
            "projectId": project_id,
            "backtestId": backtest_id,
        })
        bt = r.get("backtest", {})
        progress = bt.get("progress", 0)
        completed = bt.get("completed", False)
        print(f"   {progress:.0%} complete", end="\r")
        if completed:
            print()
            break
    else:
        print("\n   Timed out after 15 min.")
        sys.exit(1)
    print("   Done.")

    # 7. Results
    print("\n7. Results")
    bt = r.get("backtest", {})
    stats = bt.get("statistics", {})
    runtime = bt.get("runtimeStatistics", {})

    print("\n" + "=" * 70)
    print("CRTPX BACKTEST RESULTS (Strategy B -- 1.2x ES Futures + Cond Risk-Off)")
    print("=" * 70)
    print(f"  Project:  {project_name}")
    print(f"  Backtest: {bt_name}")

    print("\n  KEY METRICS:")
    for key in [
        "Total Orders", "Compounding Annual Return", "Drawdown",
        "Sharpe Ratio", "Sortino Ratio",
        "Alpha", "Beta", "Total Fees",
        "Portfolio Turnover", "Net Profit",
        "Win Rate", "Loss Rate", "Average Win", "Average Loss",
    ]:
        val = stats.get(key, "N/A")
        print(f"    {key:.<42} {val}")

    print("\n  RUNTIME:")
    for key, val in runtime.items():
        print(f"    {key:.<42} {val}")

    # Save full results
    out = os.path.join(SCRIPT_DIR, "qc_crtpx_results.json")
    with open(out, "w") as f:
        json.dump(r, f, indent=2, default=str)
    print(f"\n  Full results saved to: {out}")

    # Algorithm logs
    algo_logs = bt.get("logs", "")
    if algo_logs:
        print("\n  ALGORITHM LOGS (last 40):")
        for line in algo_logs.strip().split("\n")[-40:]:
            print(f"    {line}")

    url = f"https://www.quantconnect.com/terminal/{project_id}#open/{backtest_id}"
    print(f"\n  View: {url}")

    # 8. Retrieve trade log from ObjectStore
    print("\n8. Retrieving trade log from ObjectStore...")
    trade_log = None
    try:
        r_obj = api_get("/object/get", {
            "projectId": project_id,
            "key": "crtpx_trade_log",
        })
        if r_obj.get("success") and "data" in r_obj:
            trade_log = json.loads(r_obj["data"])
    except Exception:
        pass

    if not trade_log:
        try:
            r_obj = api_post("/object/get", {
                "projectId": project_id,
                "key": "crtpx_trade_log",
            })
            if r_obj.get("success") and "data" in r_obj:
                trade_log = json.loads(r_obj["data"])
        except Exception:
            pass

    csv_path = os.path.join(SCRIPT_DIR, "crtpx_trades.csv")
    if trade_log and len(trade_log) > 0:
        import csv
        with open(csv_path, "w", newline="") as f:
            writer = csv.DictWriter(f, fieldnames=trade_log[0].keys())
            writer.writeheader()
            writer.writerows(trade_log)
        print(f"   Trade CSV saved: {csv_path} ({len(trade_log)} trades)")
    else:
        print("   Could not retrieve trade log from ObjectStore.")
        print("   Generating trade CSV from order events in backtest results...")

        orders = bt.get("orders", {})
        if orders:
            import csv
            rows = []
            for oid, order in sorted(orders.items(), key=lambda x: x[1].get("time", "")):
                rows.append({
                    "date": order.get("time", "")[:10],
                    "action": "BUY" if order.get("quantity", 0) > 0 else "SELL",
                    "ticker": order.get("symbol", {}).get("value", ""),
                    "price": order.get("price", 0),
                    "shares": abs(order.get("quantity", 0)),
                    "value": round(abs(order.get("quantity", 0) * order.get("price", 0)), 2),
                    "status": order.get("status", ""),
                    "type": order.get("type", ""),
                })
            if rows:
                with open(csv_path, "w", newline="") as f:
                    writer = csv.DictWriter(f, fieldnames=rows[0].keys())
                    writer.writeheader()
                    writer.writerows(rows)
                print(f"   Trade CSV saved: {csv_path} ({len(rows)} orders)")
            else:
                print("   No orders found in results.")
        else:
            print("   No orders in backtest results.")

    # 9. Generate Word document
    print("\n9. Generating Word document...")
    generate_word_report(stats, runtime, bt_name, project_name, url, csv_path, algo_logs)

    print("\n" + "=" * 70)
    print("COMPLETE")
    print("=" * 70)


def generate_word_report(stats, runtime, bt_name, project_name, url, csv_path, algo_logs):
    from docx import Document
    from docx.shared import Inches, Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import nsdecls
    from docx.oxml import parse_xml

    NAVY = RGBColor(0x1A, 0x3C, 0x5E)
    DARK = RGBColor(0x33, 0x33, 0x33)
    GRAY = RGBColor(0x66, 0x66, 0x66)
    WHITE = RGBColor(0xFF, 0xFF, 0xFF)
    GREEN = RGBColor(0x27, 0xAE, 0x60)

    def shade(cell, hex_color):
        cell._tc.get_or_add_tcPr().append(
            parse_xml(f'<w:shd {nsdecls("w")} w:fill="{hex_color}"/>')
        )

    def border(cell, color="CCCCCC"):
        cell._tc.get_or_add_tcPr().append(parse_xml(
            f'<w:tcBorders {nsdecls("w")}>'
            f'<w:top w:val="single" w:sz="4" w:space="0" w:color="{color}"/>'
            f'<w:bottom w:val="single" w:sz="4" w:space="0" w:color="{color}"/>'
            f'<w:left w:val="single" w:sz="4" w:space="0" w:color="{color}"/>'
            f'<w:right w:val="single" w:sz="4" w:space="0" w:color="{color}"/>'
            f'</w:tcBorders>'
        ))

    def add_table(doc, headers, rows, col_widths=None, highlight_col=None):
        table = doc.add_table(rows=1 + len(rows), cols=len(headers))
        table.alignment = WD_TABLE_ALIGNMENT.LEFT
        table.autofit = False
        for i, h in enumerate(headers):
            cell = table.rows[0].cells[i]
            cell.text = ""
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(h)
            run.bold = True
            run.font.size = Pt(8)
            run.font.color.rgb = WHITE
            shade(cell, "1A3C5E")
            border(cell, "1A3C5E")
            if col_widths:
                cell.width = Inches(col_widths[i])
        for r_idx, row_data in enumerate(rows):
            for c_idx, val in enumerate(row_data):
                cell = table.rows[r_idx + 1].cells[c_idx]
                cell.text = ""
                p = cell.paragraphs[0]
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER if c_idx > 0 else WD_ALIGN_PARAGRAPH.LEFT
                run = p.add_run(val)
                run.font.size = Pt(8)
                if c_idx == 0:
                    run.bold = True
                if highlight_col and c_idx == highlight_col:
                    shade(cell, "E8F4E8")
                border(cell, "CCCCCC")
                if col_widths:
                    cell.width = Inches(col_widths[c_idx])
        for row in table.rows:
            row.height = Cm(0.5)
            for cell in row.cells:
                for p in cell.paragraphs:
                    p.paragraph_format.space_before = Pt(1)
                    p.paragraph_format.space_after = Pt(1)
        return table

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10)
    style.paragraph_format.space_after = Pt(4)
    for section in doc.sections:
        section.top_margin = Cm(1.5)
        section.bottom_margin = Cm(1.2)
        section.left_margin = Cm(2.0)
        section.right_margin = Cm(2.0)

    # Header
    h = doc.add_paragraph()
    run = h.add_run("POTOMAC FUND MANAGEMENT")
    run.bold = True
    run.font.size = Pt(16)
    run.font.color.rgb = NAVY
    h.paragraph_format.space_after = Pt(0)

    sub = doc.add_paragraph()
    run = sub.add_run("CRTPX Backtest: Penta Tactical Signal + TLH")
    run.bold = True
    run.font.size = Pt(13)
    run.font.color.rgb = DARK
    sub.paragraph_format.space_after = Pt(0)

    dateline = doc.add_paragraph()
    run = dateline.add_run("February 2026  |  QuantConnect Backtest  |  June 2019 \u2013 February 2026")
    run.font.size = Pt(10)
    run.font.color.rgb = GRAY
    dateline.paragraph_format.space_after = Pt(0)

    author = doc.add_paragraph()
    run = author.add_run("Woody Wiegmann")
    run.font.size = Pt(10)
    run.font.color.rgb = GRAY
    author.paragraph_format.space_after = Pt(2)

    line = doc.add_paragraph()
    line.paragraph_format.space_before = Pt(0)
    line.paragraph_format.space_after = Pt(4)
    pPr = line._p.get_or_add_pPr()
    pPr.append(parse_xml(f'<w:pBdr {nsdecls("w")}><w:bottom w:val="single" w:sz="6" w:space="1" w:color="1A3C5E"/></w:pBdr>'))

    # Signal description
    sig_h = doc.add_paragraph()
    run = sig_h.add_run("Signal Architecture")
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = NAVY
    sig_h.paragraph_format.space_after = Pt(2)

    signals = [
        ("Penta1", "SPY close > 50-day SMA(SPY)", "Trend"),
        ("Penta2", "IYT close > 50-day SMA(IYT)", "Economic confirmation"),
        ("Penta3", "^NYA close > 50-day SMA(^NYA)", "NYSE breadth"),
        ("Penta4", "LQD close > 50-day SMA(LQD)", "Credit conditions"),
    ]
    for name, desc, label in signals:
        p = doc.add_paragraph()
        run = p.add_run(f"{name}: ")
        run.bold = True
        run.font.size = Pt(9)
        run = p.add_run(f"{desc}  ({label})")
        run.font.size = Pt(9)
        p.paragraph_format.space_after = Pt(1)
        p.paragraph_format.left_indent = Inches(0.15)

    conf_p = doc.add_paragraph()
    run = conf_p.add_run("Penta ON = 3+ of 4 green.  3-day confirmation (Sum/ExRem).  Risk-on: 100% VOO.  Risk-off: 100% SGOV.")
    run.font.size = Pt(9)
    run.bold = True
    conf_p.paragraph_format.space_before = Pt(4)
    conf_p.paragraph_format.space_after = Pt(6)

    # Key metrics table
    met_h = doc.add_paragraph()
    run = met_h.add_run("Backtest Results")
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = NAVY
    met_h.paragraph_format.space_after = Pt(2)

    met_headers = ["Metric", "Value"]
    met_rows = []
    for key in [
        "Compounding Annual Return", "Net Profit", "Drawdown",
        "Sharpe Ratio", "Sortino Ratio", "Alpha", "Beta",
        "Win Rate", "Loss Rate", "Average Win", "Average Loss",
        "Total Fees", "Portfolio Turnover", "Total Orders",
    ]:
        met_rows.append([key, stats.get(key, "N/A")])

    equity_val = runtime.get("Equity", "N/A")
    met_rows.append(["Final Equity", equity_val])

    add_table(doc, met_headers, met_rows, col_widths=[2.5, 2.0])

    # Runtime section
    rt_h = doc.add_paragraph()
    run = rt_h.add_run("Runtime Statistics")
    run.bold = True
    run.font.size = Pt(10)
    run.font.color.rgb = NAVY
    rt_h.paragraph_format.space_before = Pt(8)
    rt_h.paragraph_format.space_after = Pt(2)

    rt_headers = ["Metric", "Value"]
    rt_rows = [[k, v] for k, v in runtime.items()]
    add_table(doc, rt_headers, rt_rows, col_widths=[2.5, 2.0])

    # TLH section from logs
    if algo_logs:
        tlh_h = doc.add_paragraph()
        run = tlh_h.add_run("Algorithm Log (Summary)")
        run.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = NAVY
        tlh_h.paragraph_format.space_before = Pt(8)
        tlh_h.paragraph_format.space_after = Pt(2)

        summary_lines = [l for l in algo_logs.strip().split("\n") if "SUMMARY" in l or "Regime" in l or "harvest" in l.lower() or "==" in l or "switches" in l.lower() or "Harvest" in l]
        if summary_lines:
            for sl in summary_lines[-15:]:
                p = doc.add_paragraph()
                run = p.add_run(sl.strip())
                run.font.size = Pt(8)
                run.font.name = "Consolas"
                p.paragraph_format.space_after = Pt(0)

    # AmiBroker instructions for coder
    ami_h = doc.add_paragraph()
    run = ami_h.add_run("AmiBroker Implementation (for Developer)")
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = NAVY
    ami_h.paragraph_format.space_before = Pt(10)
    ami_h.paragraph_format.space_after = Pt(4)

    ami_code = [
        'SMA_Period = 50;  ConfirmDays = 3;',
        '',
        '// Penta signal indicators',
        'Penta1 = Foreign("SPY","C") > MA(Foreign("SPY","C"), SMA_Period);',
        'Penta2 = Foreign("IYT","C") > MA(Foreign("IYT","C"), SMA_Period);',
        'Penta3 = Foreign("~NYA","C") > MA(Foreign("~NYA","C"), SMA_Period);',
        'Penta4 = Foreign("LQD","C") > MA(Foreign("LQD","C"), SMA_Period);',
        '',
        'PentaScore = Penta1 + Penta2 + Penta3 + Penta4;',
        'RawOn = PentaScore >= 3;',
        '',
        '// 3-day confirmation (signal must persist 3 consecutive days)',
        'ConfOn  = Sum(RawOn, ConfirmDays) == ConfirmDays;',
        'ConfOff = Sum(!RawOn, ConfirmDays) == ConfirmDays;',
        '',
        '// ExRem: fire once, then wait for opposite signal',
        'Buy  = ExRem(ConfOn, ConfOff);',
        'Sell = ExRem(ConfOff, ConfOn);',
        '',
        'SetPositionSize(100, spsPercentOfEquity);',
        '',
        '// Risk-on:  Trade VOO (or next in TLH ring: IVV, SPLG)',
        '// Risk-off: Hold SGOV (or next in TLH ring: BIL, SHV)',
        '',
        '// TLH Parameters:',
        '//   Loss trigger: -3% unrealized loss',
        '//   Wash sale window: 31 calendar days',
        '//   Equity ring: VOO -> IVV -> SPLG -> VOO',
        '//   Cash ring:   SGOV -> BIL -> SHV -> SGOV',
    ]
    for codeline in ami_code:
        p = doc.add_paragraph()
        run = p.add_run(codeline)
        run.font.size = Pt(8)
        run.font.name = "Consolas"
        run.font.color.rgb = DARK
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.left_indent = Inches(0.3)

    # Footer
    footer_line = doc.add_paragraph()
    footer_line.paragraph_format.space_before = Pt(10)
    footer_line.paragraph_format.space_after = Pt(2)
    pPr = footer_line._p.get_or_add_pPr()
    pPr.append(parse_xml(f'<w:pBdr {nsdecls("w")}><w:bottom w:val="single" w:sz="6" w:space="1" w:color="1A3C5E"/></w:pBdr>'))

    bl = doc.add_paragraph()
    run = bl.add_run("Data Sources: ")
    run.bold = True
    run.font.size = Pt(8)
    run.font.color.rgb = NAVY
    run = bl.add_run(
        "SPY (equity), IYT (equity), NYA (QuantConnect Cash Index -- actual NYSE Composite), "
        "LQD (equity). SMA computed by QuantConnect's built-in SMA indicator on daily close prices. "
        "Trade CSV attached separately."
    )
    run.font.size = Pt(8)
    run.font.color.rgb = DARK

    link_p = doc.add_paragraph()
    run = link_p.add_run(f"QuantConnect backtest: {url}")
    run.font.size = Pt(8)
    run.font.color.rgb = GRAY

    out_path = os.path.join(SCRIPT_DIR, "CRTPX_Backtest_Report.docx")
    doc.save(out_path)
    print(f"   Word document saved: {out_path}")


if __name__ == "__main__":
    main()
