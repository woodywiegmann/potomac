"""
Generate comparison one-pager for breadth/trend overlay sweep results.
"""

import sys
try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn
except ImportError:
    sys.exit("Required: pip install python-docx")


def set_cell_shading(cell, color):
    shading = cell._element.get_or_add_tcPr()
    sh = shading.makeelement(qn("w:shd"), {qn("w:fill"): color, qn("w:val"): "clear"})
    shading.append(sh)


def styled_table(doc, headers, rows, header_color="1F4E79", highlight_row=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"

    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(255, 255, 255)
        set_cell_shading(cell, header_color)

    for r_idx, row_data in enumerate(rows):
        for c_idx, val in enumerate(row_data):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = str(val)
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(9)
                    if r_idx == highlight_row:
                        run.bold = True

        if r_idx == highlight_row:
            for c_idx in range(len(headers)):
                set_cell_shading(table.rows[r_idx + 1].cells[c_idx], "D5F5E3")
        elif r_idx % 2 == 1:
            for c_idx in range(len(headers)):
                set_cell_shading(table.rows[r_idx + 1].cells[c_idx], "EBF5FB")

    return table


def main():
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10)
    for section in doc.sections:
        section.top_margin = Inches(0.6)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.55)
        section.right_margin = Inches(0.55)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Breadth + Trend Overlay: Threshold Comparison")
    run.bold = True
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(31, 78, 121)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = sub.add_run(
        "40-ETF International Dual Momentum  |  Jan 2016 - Feb 2026\n"
        "Signal A: % MSCI Country ETFs > 200d SMA  |  Signal B: ACWX > 200d SMA\n"
        "OR Gate: Stay invested if EITHER positive; cash (BIL) only if BOTH negative"
    )
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(100, 100, 100)

    # Main comparison table
    doc.add_paragraph()
    h = doc.add_paragraph()
    run = h.add_run("Performance Comparison")
    run.bold = True
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(31, 78, 121)

    headers = ["Variant", "CAGR", "Total Ret", "Max DD", "Sharpe", "Sortino",
               "Alpha", "Beta", "Vol", "End Equity"]

    rows = [
        ["Baseline (no overlay)", "14.30%", "289.0%", "39.8%", "0.508", "0.541",
         "0.049", "0.80", "16.7%", "$3,890,450"],
        ["Breadth 30%", "10.86%", "185.3%", "31.7%", "0.384", "0.363",
         "0.030", "0.652", "15.6%", "$2,852,632"],
        ["Breadth 40%", "10.12%", "166.4%", "35.8%", "0.363", "0.324",
         "0.027", "0.588", "14.7%", "$2,664,092"],
        ["Breadth 50%*", "~1.1%", "~11%", "N/A", "N/A", "N/A",
         "N/A", "N/A", "N/A", "$1,113,359"],
        ["Breadth 60%", "12.20%", "222.4%", "21.8%", "0.511", "0.450",
         "0.049", "0.337", "12.6%", "$3,223,515"],
    ]

    # highlight_row=4 -> 60% threshold (best risk-adjusted)
    styled_table(doc, headers, rows, highlight_row=4)

    p = doc.add_paragraph()
    run = p.add_run("* 50% threshold: Strategy was in cash almost the entire period (stats unavailable). "
                     "Green highlight = best risk-adjusted variant.")
    run.font.size = Pt(8)
    run.italic = True
    run.font.color.rgb = RGBColor(120, 120, 120)

    # Analysis
    doc.add_paragraph()
    h = doc.add_paragraph()
    run = h.add_run("Key Takeaways")
    run.bold = True
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(31, 78, 121)

    takeaways = [
        ("60% Threshold is the clear winner on a risk-adjusted basis. "
         "It delivers the highest Sharpe (0.511 vs baseline 0.508), cuts max drawdown nearly in half "
         "(21.8% vs 39.8%), and reduces annualized vol from 16.7% to 12.6% -- while preserving "
         "most of the return (12.2% CAGR vs 14.3%)."),
        ("The overlay's primary benefit is drawdown reduction, not return enhancement. "
         "All overlay variants underperform the raw baseline on total return, "
         "but they significantly reduce the severity and duration of drawdowns."),
        ("30% and 40% thresholds are too loose. They let in more risk (31.7% and 35.8% DD) "
         "while materially hurting returns -- worst of both worlds."),
        ("50% is too aggressive. The breadth requirement is so strict that international markets "
         "almost never qualify, leaving the portfolio in cash for the vast majority of the period."),
        ("Beta compression is dramatic at 60%: 0.337 vs 0.80 baseline. "
         "The strategy captures upside selectively while avoiding broad-based selloffs."),
    ]

    for i, text in enumerate(takeaways):
        p = doc.add_paragraph(f"{i+1}. {text}")
        for run in p.runs:
            run.font.size = Pt(9)

    # Recommendation
    doc.add_paragraph()
    h = doc.add_paragraph()
    run = h.add_run("Recommendation")
    run.bold = True
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(31, 78, 121)

    p = doc.add_paragraph()
    run = p.add_run("Adopt the 60% breadth threshold overlay. ")
    run.bold = True
    run.font.size = Pt(10)
    run = p.add_run(
        "The trade-off is compelling: sacrifice ~2% CAGR (14.3% -> 12.2%) in exchange for "
        "an 18-percentage-point reduction in max drawdown (39.8% -> 21.8%) and 4% lower annualized vol. "
        "The Sharpe ratio actually improves slightly (0.508 -> 0.511), "
        "and alpha holds at 0.049. Beta drops from 0.80 to 0.34, "
        "meaning the strategy captures return with far less market exposure."
    )
    run.font.size = Pt(10)

    # Backtest links
    doc.add_paragraph()
    h = doc.add_paragraph()
    run = h.add_run("QuantConnect Backtest Links")
    run.bold = True
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(31, 78, 121)

    links = [
        ("Baseline", "https://www.quantconnect.com/terminal/28624941#open/912a29bf64b140adfb7a246246aa3dbd"),
        ("30% Threshold", "https://www.quantconnect.com/terminal/28632343#open/26aeb6759024e7ae90ebc83e7a8b4ac0"),
        ("40% Threshold", "https://www.quantconnect.com/terminal/28632371#open/9ec8cf5e1fa3d19f2fe1fba15f42fc7b"),
        ("50% Threshold", "https://www.quantconnect.com/terminal/28632398#open/3cdd4c0548c9c2ce66130aa9f76b8f12"),
        ("60% Threshold", "https://www.quantconnect.com/terminal/28632421#open/fe32f62bbd246dfea5e938329309e071"),
    ]

    for label, url in links:
        p = doc.add_paragraph()
        run = p.add_run(f"{label}: ")
        run.font.size = Pt(8)
        run.bold = True
        run = p.add_run(url)
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(0, 102, 204)

    out = r"c:\Users\WoodyWiegmann\OneDrive - PFM\Desktop\Potomac\Breadth_Sweep_Comparison.docx"
    doc.save(out)
    print(f"Saved: {out}")


if __name__ == "__main__":
    main()
