"""
Save 'BEST IDEA SO FAR' -- 40-ETF International Dual Momentum baseline results.
"""

import sys
try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
except ImportError:
    sys.exit("Required: pip install python-docx")


def set_cell_shading(cell, color):
    shading = cell._element.get_or_add_tcPr()
    sh = shading.makeelement(qn("w:shd"), {qn("w:fill"): color, qn("w:val"): "clear"})
    shading.append(sh)


def styled_table(doc, headers, rows, header_color="1F4E79"):
    from docx.enum.table import WD_TABLE_ALIGNMENT
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(255, 255, 255)
        set_cell_shading(cell, header_color)
    for r_idx, row_data in enumerate(rows):
        for c_idx, val in enumerate(row_data):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = str(val)
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(9)
        if r_idx % 2 == 1:
            for c_idx in range(len(headers)):
                set_cell_shading(table.rows[r_idx + 1].cells[c_idx], "EBF5FB")
    return table


def main():
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10)
    for section in doc.sections:
        section.top_margin = Inches(0.6)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.6)
        section.right_margin = Inches(0.6)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("BEST IDEA SO FAR")
    run.bold = True
    run.font.size = Pt(18)
    run.font.color.rgb = RGBColor(31, 78, 121)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = sub.add_run("International 40-ETF Blended-Lookback Dual Momentum\nPotomac Fund Management  |  March 2026")
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(100, 100, 100)

    # Strategy summary
    doc.add_paragraph()
    h = doc.add_paragraph()
    run = h.add_run("Strategy Summary")
    run.bold = True
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(31, 78, 121)

    params = [
        ("Universe", "40 international ETFs (24 developed, 16 EM) -- optimized for minimum pairwise correlation"),
        ("Signal", "Blended momentum: average of 1-month, 3-month, 6-month, and 12-month trailing returns"),
        ("Ranking", "All 40 ETFs ranked by composite score each month"),
        ("Holdings", "Top 7, equal-weight (~14.3% each)"),
        ("Absolute Momentum", "Composite score must be > 0 to be held; failing slots go to BIL"),
        ("Risk-Off", "BIL (SPDR Bloomberg 1-3 Month T-Bill ETF)"),
        ("Rebalance", "Monthly (last trading day)"),
        ("Benchmark", "EFA (iShares MSCI EAFE)"),
    ]
    styled_table(doc, ["Parameter", "Detail"], params)

    # Backtest results
    doc.add_paragraph()
    h = doc.add_paragraph()
    run = h.add_run("QuantConnect Backtest Results  (Jan 2016 - Feb 2026)")
    run.bold = True
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(31, 78, 121)

    metrics = [
        ("Compounding Annual Return", "14.30%"),
        ("Total Return (Net Profit)", "289.0%"),
        ("Max Drawdown", "39.8%"),
        ("Sharpe Ratio", "0.508"),
        ("Sortino Ratio", "0.541"),
        ("Alpha", "+4.9%"),
        ("Beta", "0.80"),
        ("Annual Standard Deviation", "16.7%"),
        ("Information Ratio", "0.314"),
        ("Tracking Error", "12.5%"),
        ("Win Rate", "58%"),
        ("Profit-Loss Ratio", "1.33"),
        ("Total Orders", "486"),
        ("Total Fees", "$12,580"),
        ("Starting Equity", "$1,000,000"),
        ("Ending Equity", "$3,890,450"),
        ("Portfolio Turnover", "1.88%"),
        ("Estimated Capacity", "$6.3M"),
        ("Lowest Capacity Asset", "LIT"),
    ]
    styled_table(doc, ["Metric", "Value"], metrics)

    # Universe composition
    doc.add_paragraph()
    h = doc.add_paragraph()
    run = h.add_run("Universe Composition (40 ETFs)")
    run.bold = True
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(31, 78, 121)

    universe = [
        ("Dev Country (10)", "EWJ, EWG, EWQ, EWI, EWD, EWL, EWP, EWH, EWS, EDEN"),
        ("Dev Factor (1)", "IHDG"),
        ("Dev Thematic (13)", "RING, SIL, URA, KXI, LIT, REMX, COPX, PICK, GNR, CGW, GII, INFL, MOO"),
        ("EM Country (14)", "EWT, EWZ, INDA, FXI, EWY, EWW, ILF, ECH, TUR, ARGT, VNM, THD, EWM, EIDO"),
        ("EM Broad (2)", "KSA, KWEB"),
    ]
    styled_table(doc, ["Bucket", "Tickers"], universe, header_color="2C6E49")

    # Correlation stats
    doc.add_paragraph()
    h = doc.add_paragraph()
    run = h.add_run("Correlation Optimization")
    run.bold = True
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(31, 78, 121)

    corr = [
        ("Full candidate universe (71 ETFs)", "0.5632"),
        ("Trimmed universe (40 ETFs)", "0.4750"),
        ("Correlation reduction", "-15.7%"),
        ("Pairs > 0.90 (full 71)", "65"),
        ("Pairs > 0.90 (trimmed 40)", "7"),
        ("Dev / EM split", "60% / 40%"),
    ]
    styled_table(doc, ["Metric", "Value"], corr)

    # Link
    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run("QuantConnect Backtest: ")
    run.font.size = Pt(9)
    run = p.add_run("https://www.quantconnect.com/terminal/28624941#open/912a29bf64b140adfb7a246246aa3dbd")
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0, 102, 204)

    p = doc.add_paragraph()
    run = p.add_run("Project ID: 28624941  |  Backtest ID: 912a29bf64b140adfb7a246246aa3dbd")
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(150, 150, 150)

    out = r"c:\Users\WoodyWiegmann\OneDrive - PFM\Desktop\Potomac\BEST_IDEA_SO_FAR.docx"
    doc.save(out)
    print(f"Saved: {out}")


if __name__ == "__main__":
    main()
