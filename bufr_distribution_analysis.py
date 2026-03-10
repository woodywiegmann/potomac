"""
BUFR (FT Vest Laddered Buffer ETF) -- Return Distribution & Capture Analysis
==============================================================================
Largest buffer ETF by AUM (~$8.5B). Analyzes:
  1. Daily return distribution vs SPY -- skew, kurtosis, tail exposure
  2. Upside / downside capture ratios (monthly, Morningstar methodology)
  3. Left-tail vs right-tail asymmetry
  4. Generates a one-pager Word document with results
"""

import math, os, sys, warnings
warnings.filterwarnings("ignore")

try:
    import yfinance as yf
    import pandas as pd
    import numpy as np
    from scipy import stats as sp_stats
except ImportError:
    print("Required: pip install yfinance pandas numpy scipy")
    sys.exit(1)

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml

SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
OUT_DIR = r"c:\Users\WoodyWiegmann\OneDrive - PFM\Desktop\Portfolio Ideas"

NAVY  = RGBColor(0x1A, 0x3C, 0x5E)
DARK  = RGBColor(0x33, 0x33, 0x33)
GRAY  = RGBColor(0x66, 0x66, 0x66)
RED   = RGBColor(0xC0, 0x39, 0x2B)
GREEN = RGBColor(0x1E, 0x7D, 0x32)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)


# ═══════════════════════════════════════════════════════════════════════════════
# DATA
# ═══════════════════════════════════════════════════════════════════════════════

def fetch(tickers, start):
    data = yf.download(tickers, start=start, auto_adjust=True, progress=False)
    if isinstance(data.columns, pd.MultiIndex):
        return data["Close"].ffill().dropna(how="all")
    return data[["Close"]].rename(columns={"Close": tickers[0]}).ffill()


# ═══════════════════════════════════════════════════════════════════════════════
# ANALYTICS
# ═══════════════════════════════════════════════════════════════════════════════

def daily_stats(daily_returns, label):
    return {
        "label": label,
        "mean": daily_returns.mean() * 252,
        "vol": daily_returns.std() * math.sqrt(252),
        "skew": daily_returns.skew(),
        "kurt": daily_returns.kurtosis(),
        "min": daily_returns.min(),
        "max": daily_returns.max(),
        "median": daily_returns.median(),
        "pct_positive": (daily_returns > 0).mean(),
    }


def tail_analysis(daily_ret, benchmark_ret, percentiles=[1, 5, 10]):
    results = {}
    for p in percentiles:
        threshold = np.percentile(benchmark_ret.dropna(), p)
        bench_tail_days = benchmark_ret[benchmark_ret <= threshold]
        fund_on_those_days = daily_ret.reindex(bench_tail_days.index).dropna()
        if len(fund_on_those_days) == 0:
            continue

        bench_avg = bench_tail_days.reindex(fund_on_those_days.index).mean()
        fund_avg = fund_on_those_days.mean()
        capture = fund_avg / bench_avg if bench_avg != 0 else 0

        results[p] = {
            "n_days": len(fund_on_those_days),
            "bench_avg": bench_avg,
            "fund_avg": fund_avg,
            "capture_pct": capture * 100,
            "protection_pct": (1 - capture) * 100,
        }
    return results


def upside_tail_analysis(daily_ret, benchmark_ret, percentiles=[90, 95, 99]):
    results = {}
    for p in percentiles:
        threshold = np.percentile(benchmark_ret.dropna(), p)
        bench_tail_days = benchmark_ret[benchmark_ret >= threshold]
        fund_on_those_days = daily_ret.reindex(bench_tail_days.index).dropna()
        if len(fund_on_those_days) == 0:
            continue

        bench_avg = bench_tail_days.reindex(fund_on_those_days.index).mean()
        fund_avg = fund_on_those_days.mean()
        capture = fund_avg / bench_avg if bench_avg != 0 else 0

        results[p] = {
            "n_days": len(fund_on_those_days),
            "bench_avg": bench_avg,
            "fund_avg": fund_avg,
            "capture_pct": capture * 100,
        }
    return results


def monthly_capture(fund_ret, bench_ret):
    """Morningstar-style monthly upside/downside capture ratios."""
    fund_monthly = (1 + fund_ret).resample("ME").prod() - 1
    bench_monthly = (1 + bench_ret).resample("ME").prod() - 1

    merged = pd.DataFrame({"fund": fund_monthly, "bench": bench_monthly}).dropna()

    up_months = merged[merged["bench"] > 0]
    down_months = merged[merged["bench"] < 0]

    if len(up_months) > 0:
        up_geo_fund = (1 + up_months["fund"]).prod() ** (1/len(up_months)) - 1
        up_geo_bench = (1 + up_months["bench"]).prod() ** (1/len(up_months)) - 1
        upside_capture = (up_geo_fund / up_geo_bench) * 100 if up_geo_bench != 0 else 0
    else:
        upside_capture = 0

    if len(down_months) > 0:
        down_geo_fund = (1 + down_months["fund"]).prod() ** (1/len(down_months)) - 1
        down_geo_bench = (1 + down_months["bench"]).prod() ** (1/len(down_months)) - 1
        downside_capture = (down_geo_fund / down_geo_bench) * 100 if down_geo_bench != 0 else 0
    else:
        downside_capture = 0

    return {
        "upside": upside_capture,
        "downside": downside_capture,
        "up_months": len(up_months),
        "down_months": len(down_months),
        "up_avg_fund": up_months["fund"].mean() if len(up_months) > 0 else 0,
        "up_avg_bench": up_months["bench"].mean() if len(up_months) > 0 else 0,
        "down_avg_fund": down_months["fund"].mean() if len(down_months) > 0 else 0,
        "down_avg_bench": down_months["bench"].mean() if len(down_months) > 0 else 0,
    }


def capture_by_period(fund_ret, bench_ret):
    """Capture ratios for 1Y, 3Y, since inception."""
    results = {}
    now = fund_ret.index[-1]
    periods = {
        "1-Year": fund_ret.index >= now - pd.DateOffset(years=1),
        "3-Year": fund_ret.index >= now - pd.DateOffset(years=3),
        "Since Inception": fund_ret.index >= fund_ret.index[0],
    }
    for label, mask in periods.items():
        f = fund_ret[mask]
        b = bench_ret.reindex(f.index).dropna()
        f = f.reindex(b.index)
        if len(f) < 30:
            continue
        results[label] = monthly_capture(f, b)
    return results


# ═══════════════════════════════════════════════════════════════════════════════
# WORD DOC HELPERS
# ═══════════════════════════════════════════════════════════════════════════════

def set_cell_shading(cell, hex_color):
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{hex_color}"/>')
    cell._tc.get_or_add_tcPr().append(shading)

def set_cell_border(cell, color="CCCCCC", size="4"):
    tcPr = cell._tc.get_or_add_tcPr()
    borders = parse_xml(
        f'<w:tcBorders {nsdecls("w")}>'
        f'  <w:top w:val="single" w:sz="{size}" w:space="0" w:color="{color}"/>'
        f'  <w:bottom w:val="single" w:sz="{size}" w:space="0" w:color="{color}"/>'
        f'  <w:left w:val="single" w:sz="{size}" w:space="0" w:color="{color}"/>'
        f'  <w:right w:val="single" w:sz="{size}" w:space="0" w:color="{color}"/>'
        f'</w:tcBorders>'
    )
    tcPr.append(borders)

def add_table(doc, headers, rows, col_widths=None, highlight_rows=None,
              highlight_color="E8F4E8", red_rows=None, red_color="FDEDED"):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        run.bold = True
        run.font.size = Pt(7.5)
        run.font.color.rgb = WHITE
        set_cell_shading(cell, "1A3C5E")
        set_cell_border(cell, color="1A3C5E")
        if col_widths:
            cell.width = Inches(col_widths[i])
    for r_idx, row_data in enumerate(rows):
        for c_idx, val in enumerate(row_data):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = ""
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if c_idx > 0 else WD_ALIGN_PARAGRAPH.LEFT
            run = p.add_run(str(val))
            run.font.size = Pt(7.5)
            if c_idx == 0:
                run.bold = True
            if highlight_rows and r_idx in highlight_rows:
                set_cell_shading(cell, highlight_color)
            if red_rows and r_idx in red_rows:
                set_cell_shading(cell, red_color)
            set_cell_border(cell, color="CCCCCC")
            if col_widths:
                cell.width = Inches(col_widths[c_idx])
    for row in table.rows:
        row.height = Cm(0.45)
        for cell in row.cells:
            for p in cell.paragraphs:
                p.paragraph_format.space_before = Pt(1)
                p.paragraph_format.space_after = Pt(1)
    return table

def add_navy_line(doc):
    line = doc.add_paragraph()
    line.paragraph_format.space_before = Pt(0)
    line.paragraph_format.space_after = Pt(2)
    pPr = line._p.get_or_add_pPr()
    pBdr = parse_xml(f'<w:pBdr {nsdecls("w")}><w:bottom w:val="single" w:sz="6" w:space="1" w:color="1A3C5E"/></w:pBdr>')
    pPr.append(pBdr)

def add_section_header(doc, text):
    h = doc.add_paragraph()
    run = h.add_run(text)
    run.bold = True
    run.font.size = Pt(9.5)
    run.font.color.rgb = NAVY
    h.paragraph_format.space_before = Pt(6)
    h.paragraph_format.space_after = Pt(2)

def add_body(doc, text, bold_prefix=None, size=9):
    p = doc.add_paragraph()
    if bold_prefix:
        run = p.add_run(bold_prefix)
        run.bold = True
        run.font.size = Pt(size)
    run = p.add_run(text)
    run.font.size = Pt(size)
    p.paragraph_format.space_after = Pt(2)
    return p

def add_footnote(doc, text):
    fn = doc.add_paragraph()
    run = fn.add_run(text)
    run.font.size = Pt(7)
    run.font.color.rgb = GRAY
    fn.paragraph_format.space_before = Pt(1)
    fn.paragraph_format.space_after = Pt(4)


# ═══════════════════════════════════════════════════════════════════════════════
# MAIN
# ═══════════════════════════════════════════════════════════════════════════════

def main():
    print("=" * 70)
    print("  BUFR RETURN DISTRIBUTION & CAPTURE ANALYSIS")
    print("=" * 70)

    print("\nFetching data...")
    prices = fetch(["BUFR", "SPY"], "2020-08-10")
    daily_ret = prices.pct_change().dropna()
    bufr = daily_ret["BUFR"]
    spy = daily_ret["SPY"]

    aligned = pd.DataFrame({"BUFR": bufr, "SPY": spy}).dropna()
    bufr = aligned["BUFR"]
    spy = aligned["SPY"]

    start_date = aligned.index[0].strftime("%Y-%m-%d")
    end_date = aligned.index[-1].strftime("%Y-%m-%d")
    n_days = len(aligned)
    print(f"  {start_date} to {end_date} ({n_days} trading days)")

    # ── Distribution stats ───────────────────────────────────────────────────
    bufr_stats = daily_stats(bufr, "BUFR")
    spy_stats = daily_stats(spy, "SPY")

    print(f"\n  Distribution Statistics:")
    print(f"  {'':20s} {'BUFR':>10s} {'SPY':>10s}")
    print(f"  {'-'*42}")
    print(f"  {'Ann Return':20s} {bufr_stats['mean']:>+10.2%} {spy_stats['mean']:>+10.2%}")
    print(f"  {'Ann Vol':20s} {bufr_stats['vol']:>10.2%} {spy_stats['vol']:>10.2%}")
    print(f"  {'Skewness':20s} {bufr_stats['skew']:>+10.3f} {spy_stats['skew']:>+10.3f}")
    print(f"  {'Excess Kurtosis':20s} {bufr_stats['kurt']:>+10.3f} {spy_stats['kurt']:>+10.3f}")
    print(f"  {'Worst Day':20s} {bufr_stats['min']:>+10.2%} {spy_stats['min']:>+10.2%}")
    print(f"  {'Best Day':20s} {bufr_stats['max']:>+10.2%} {spy_stats['max']:>+10.2%}")
    print(f"  {'% Positive Days':20s} {bufr_stats['pct_positive']:>10.1%} {spy_stats['pct_positive']:>10.1%}")

    # ── Tail analysis ────────────────────────────────────────────────────────
    left_tail = tail_analysis(bufr, spy, [1, 5, 10])
    right_tail = upside_tail_analysis(bufr, spy, [90, 95, 99])

    print(f"\n  LEFT TAIL (worst SPY days):")
    print(f"  {'Percentile':>12s} {'Days':>6s} {'SPY Avg':>10s} {'BUFR Avg':>10s} {'Capture':>10s} {'Protected':>10s}")
    for p, d in left_tail.items():
        print(f"  {'Worst ' + str(p) + '%':>12s} {d['n_days']:>6d} {d['bench_avg']:>+10.2%} {d['fund_avg']:>+10.2%} {d['capture_pct']:>9.1f}% {d['protection_pct']:>9.1f}%")

    print(f"\n  RIGHT TAIL (best SPY days):")
    print(f"  {'Percentile':>12s} {'Days':>6s} {'SPY Avg':>10s} {'BUFR Avg':>10s} {'Capture':>10s}")
    for p, d in right_tail.items():
        print(f"  {'Top ' + str(100-p) + '%':>12s} {d['n_days']:>6d} {d['bench_avg']:>+10.2%} {d['fund_avg']:>+10.2%} {d['capture_pct']:>9.1f}%")

    # ── Monthly capture ratios ───────────────────────────────────────────────
    captures = capture_by_period(bufr, spy)

    print(f"\n  MONTHLY CAPTURE RATIOS (Morningstar methodology):")
    print(f"  {'Period':>18s} {'Upside':>10s} {'Downside':>10s} {'Up Mo':>7s} {'Down Mo':>8s}")
    for label, c in captures.items():
        print(f"  {label:>18s} {c['upside']:>9.1f}% {c['downside']:>9.1f}% {c['up_months']:>7d} {c['down_months']:>8d}")

    # ── Asymmetry summary ────────────────────────────────────────────────────
    incep = captures.get("Since Inception", {})
    up_cap = incep.get("upside", 0)
    dn_cap = incep.get("downside", 0)
    asymmetry = up_cap - dn_cap

    print(f"\n  ASYMMETRY SUMMARY:")
    print(f"  Upside capture:   {up_cap:.1f}%")
    print(f"  Downside capture: {dn_cap:.1f}%")
    print(f"  Spread:           {asymmetry:+.1f} pp")
    if asymmetry > 0:
        print(f"  --> BUFR captures more upside than downside (favorable)")
    else:
        print(f"  --> BUFR captures more downside than upside (unfavorable)")

    left_1 = left_tail.get(1, {}).get("capture_pct", 0)
    right_1 = right_tail.get(99, {}).get("capture_pct", 0)
    print(f"\n  Extreme tail capture:")
    print(f"    Worst 1% SPY days -> BUFR captures {left_1:.1f}% of the loss")
    print(f"    Best  1% SPY days -> BUFR captures {right_1:.1f}% of the gain")
    print(f"    Tail asymmetry: {right_1 - left_1:+.1f} pp")

    # ═══════════════════════════════════════════════════════════════════════════
    # WORD DOC
    # ═══════════════════════════════════════════════════════════════════════════
    print("\nGenerating Word document...")
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10)
    style.paragraph_format.space_after = Pt(4)
    style.paragraph_format.space_before = Pt(0)
    for section in doc.sections:
        section.top_margin = Cm(1.3)
        section.bottom_margin = Cm(1.0)
        section.left_margin = Cm(1.8)
        section.right_margin = Cm(1.8)

    # Header
    h = doc.add_paragraph()
    h.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = h.add_run("POTOMAC FUND MANAGEMENT")
    run.bold = True; run.font.size = Pt(15); run.font.color.rgb = NAVY
    h.paragraph_format.space_after = Pt(0)

    sub = doc.add_paragraph()
    run = sub.add_run("BUFR (FT Vest Laddered Buffer ETF): Return Distribution & Tail Analysis")
    run.bold = True; run.font.size = Pt(11.5); run.font.color.rgb = DARK
    sub.paragraph_format.space_after = Pt(0)

    dateline = doc.add_paragraph()
    run = dateline.add_run("March 2026  |  Woody Wiegmann")
    run.font.size = Pt(10); run.font.color.rgb = GRAY
    dateline.paragraph_format.space_after = Pt(0)

    add_navy_line(doc)

    # Context
    add_body(doc,
        f"BUFR is the largest buffer ETF by AUM (~$8.5B). It holds a laddered portfolio of 12 monthly "
        f"FT Vest 10% buffer ETFs on SPY, resetting one each month. The structure caps upside while "
        f"absorbing the first 10% of losses in each outcome period. Analysis below uses {n_days:,} trading "
        f"days ({start_date} to {end_date}).",
        bold_prefix="Context: ")

    # ── Distribution comparison ──────────────────────────────────────────────
    add_section_header(doc, "Daily Return Distribution: BUFR vs SPY")

    dist_headers = ["Metric", "BUFR", "SPY", "Difference"]
    dist_rows = [
        ["Annualized Return", f"{bufr_stats['mean']:+.2%}", f"{spy_stats['mean']:+.2%}",
         f"{bufr_stats['mean'] - spy_stats['mean']:+.2%}"],
        ["Annualized Volatility", f"{bufr_stats['vol']:.2%}", f"{spy_stats['vol']:.2%}",
         f"{bufr_stats['vol'] - spy_stats['vol']:+.2%}"],
        ["Skewness", f"{bufr_stats['skew']:+.3f}", f"{spy_stats['skew']:+.3f}",
         f"{bufr_stats['skew'] - spy_stats['skew']:+.3f}"],
        ["Excess Kurtosis", f"{bufr_stats['kurt']:+.3f}", f"{spy_stats['kurt']:+.3f}",
         f"{bufr_stats['kurt'] - spy_stats['kurt']:+.3f}"],
        ["Worst Day", f"{bufr_stats['min']:+.2%}", f"{spy_stats['min']:+.2%}",
         f"{bufr_stats['min'] - spy_stats['min']:+.2%}"],
        ["Best Day", f"{bufr_stats['max']:+.2%}", f"{spy_stats['max']:+.2%}",
         f"{bufr_stats['max'] - spy_stats['max']:+.2%}"],
        ["% Positive Days", f"{bufr_stats['pct_positive']:.1%}", f"{spy_stats['pct_positive']:.1%}",
         f"{bufr_stats['pct_positive'] - spy_stats['pct_positive']:+.1%}"],
        ["Beta to SPY", f"{aligned['BUFR'].cov(aligned['SPY']) / aligned['SPY'].var():.2f}",
         "1.00", ""],
        ["R-squared", f"{aligned.corr().iloc[0,1]**2:.2%}", "100%", ""],
    ]
    add_table(doc, dist_headers, dist_rows, col_widths=[1.6, 1.2, 1.2, 1.2])

    skew_note = "more negative" if bufr_stats["skew"] < spy_stats["skew"] else "less negative"
    kurt_note = "fatter" if bufr_stats["kurt"] > spy_stats["kurt"] else "thinner"
    add_body(doc,
        f"BUFR's daily returns have {skew_note} skew ({bufr_stats['skew']:+.3f} vs "
        f"{spy_stats['skew']:+.3f}) and {kurt_note} tails (excess kurtosis "
        f"{bufr_stats['kurt']:+.3f} vs {spy_stats['kurt']:+.3f}) than SPY.",
        size=8)

    # ── Left tail ────────────────────────────────────────────────────────────
    add_section_header(doc, "Left-Tail Exposure: How Much Downside Does BUFR Actually Buffer?")

    add_body(doc,
        "On SPY's worst days, how much of the loss does BUFR absorb? Lower capture = more protection.",
        size=8)

    lt_headers = ["SPY Worst...", "# Days", "SPY Avg", "BUFR Avg", "Downside\nCapture", "Loss\nProtected"]
    lt_rows = []
    for p in [1, 5, 10]:
        d = left_tail.get(p, {})
        if not d:
            continue
        lt_rows.append([
            f"Worst {p}%",
            f"{d['n_days']}",
            f"{d['bench_avg']:+.2%}",
            f"{d['fund_avg']:+.2%}",
            f"{d['capture_pct']:.1f}%",
            f"{d['protection_pct']:.1f}%",
        ])
    lt_highlight = list(range(len(lt_rows)))
    add_table(doc, lt_headers, lt_rows, col_widths=[1.0, 0.6, 0.9, 0.9, 0.9, 0.9],
              highlight_rows=lt_highlight, highlight_color="E8F4E8")

    # ── Right tail ───────────────────────────────────────────────────────────
    add_section_header(doc, "Right-Tail Exposure: How Much Upside Does BUFR Sacrifice?")

    add_body(doc,
        "On SPY's best days, how much of the gain does BUFR participate in? The cap structure limits upside.",
        size=8)

    rt_headers = ["SPY Best...", "# Days", "SPY Avg", "BUFR Avg", "Upside\nCapture", "Upside\nSacrificed"]
    rt_rows = []
    for p in [90, 95, 99]:
        d = right_tail.get(p, {})
        if not d:
            continue
        rt_rows.append([
            f"Top {100-p}%",
            f"{d['n_days']}",
            f"{d['bench_avg']:+.2%}",
            f"{d['fund_avg']:+.2%}",
            f"{d['capture_pct']:.1f}%",
            f"{100 - d['capture_pct']:.1f}%",
        ])
    add_table(doc, rt_headers, rt_rows, col_widths=[1.0, 0.6, 0.9, 0.9, 0.9, 0.9],
              red_rows=list(range(len(rt_rows))), red_color="FDEDED")

    # ── Monthly capture ──────────────────────────────────────────────────────
    add_section_header(doc, "Monthly Upside/Downside Capture (Morningstar Methodology)")

    add_body(doc,
        "Computed using geometric monthly returns. Upside capture = BUFR's geometric return in up months / "
        "SPY's geometric return in up months. Downside capture = same for down months. "
        "Ideal: high upside capture + low downside capture.",
        size=8)

    cap_headers = ["Period", "Upside\nCapture", "Downside\nCapture", "Spread\n(Up - Down)",
                   "Up\nMonths", "Down\nMonths"]
    cap_rows = []
    for label, c in captures.items():
        spread = c["upside"] - c["downside"]
        cap_rows.append([
            label,
            f"{c['upside']:.1f}%",
            f"{c['downside']:.1f}%",
            f"{spread:+.1f} pp",
            f"{c['up_months']}",
            f"{c['down_months']}",
        ])
    add_table(doc, cap_headers, cap_rows, col_widths=[1.1, 0.8, 0.8, 0.9, 0.6, 0.6])

    # ── Asymmetry verdict ────────────────────────────────────────────────────
    add_section_header(doc, "Tail Asymmetry Verdict")

    left_1_cap = left_tail.get(1, {}).get("capture_pct", 0)
    left_5_cap = left_tail.get(5, {}).get("capture_pct", 0)
    right_1_cap = right_tail.get(99, {}).get("capture_pct", 0)
    right_5_cap = right_tail.get(95, {}).get("capture_pct", 0)

    asym_headers = ["Tail", "SPY Threshold", "BUFR\nCapture", "Interpretation"]
    asym_rows = [
        ["Worst 1% days", f"{left_tail.get(1,{}).get('bench_avg',0):+.2%} avg",
         f"{left_1_cap:.1f}%",
         f"BUFR absorbs {100-left_1_cap:.0f}% of extreme left-tail losses"],
        ["Worst 5% days", f"{left_tail.get(5,{}).get('bench_avg',0):+.2%} avg",
         f"{left_5_cap:.1f}%",
         f"BUFR absorbs {100-left_5_cap:.0f}% of moderate left-tail losses"],
        ["Best 5% days", f"{right_tail.get(95,{}).get('bench_avg',0):+.2%} avg",
         f"{right_5_cap:.1f}%",
         f"BUFR sacrifices {100-right_5_cap:.0f}% of moderate right-tail gains"],
        ["Best 1% days", f"{right_tail.get(99,{}).get('bench_avg',0):+.2%} avg",
         f"{right_1_cap:.1f}%",
         f"BUFR sacrifices {100-right_1_cap:.0f}% of extreme right-tail gains"],
    ]
    add_table(doc, asym_headers, asym_rows, col_widths=[1.0, 1.0, 0.7, 3.5],
              highlight_rows=[0, 1], highlight_color="E8F4E8",
              red_rows=[2, 3], red_color="FDEDED")

    # Bottom line
    add_navy_line(doc)

    bl = doc.add_paragraph()
    run = bl.add_run("Bottom Line: ")
    run.bold = True; run.font.size = Pt(9); run.font.color.rgb = NAVY

    if up_cap > dn_cap:
        verdict = (
            f"BUFR's upside capture ({up_cap:.0f}%) exceeds its downside capture ({dn_cap:.0f}%), "
            f"a {asymmetry:+.0f} pp spread that benefits investors over full market cycles. "
        )
    else:
        verdict = (
            f"BUFR's downside capture ({dn_cap:.0f}%) exceeds its upside capture ({up_cap:.0f}%), "
            f"a {asymmetry:+.0f} pp spread that costs investors over full market cycles. "
        )

    verdict += (
        f"However, in the tails the picture sharpens: on SPY's worst 1% of days, BUFR captures "
        f"only {left_1_cap:.0f}% of the loss, but on SPY's best 1% of days it captures just "
        f"{right_1_cap:.0f}% of the gain. "
        f"The buffer structure provides meaningful crash protection but at the cost of "
        f"truncated upside in strong rallies -- a structural trade-off embedded in the options overlay."
    )
    run = bl.add_run(verdict)
    run.font.size = Pt(9); run.font.color.rgb = DARK

    add_footnote(doc,
        f"Analysis period: {start_date} to {end_date} ({n_days:,} trading days). "
        f"Monthly capture ratios use Morningstar geometric methodology. "
        f"Source: Yahoo Finance adjusted close prices. Past performance does not guarantee future results.")

    # Save
    os.makedirs(OUT_DIR, exist_ok=True)
    out1 = os.path.join(OUT_DIR, "BUFR Distribution Analysis - Mar 2026.docx")
    doc.save(out1)
    print(f"\nSaved to: {out1}")
    out2 = os.path.join(SCRIPT_DIR, "BUFR Distribution Analysis - Mar 2026.docx")
    doc.save(out2)
    print(f"Also saved to: {out2}")


if __name__ == "__main__":
    main()
