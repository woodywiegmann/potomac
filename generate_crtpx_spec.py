"""
Generate formal CRTPX strategy specification document for Dan.
Two strategies: Passive + TLH, and Enhanced (Convex Risk-Off) + TLH.
Includes AmiBroker AFL code, parameters, and academic rationale.
"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import os

SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))


def set_cell(cell, text, bold=False, size=8, color=None, align="left"):
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)
    run.font.name = "Calibri"
    if align == "center":
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif align == "right":
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT


def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0, 51, 102)


def add_para(doc, text, bold=False, size=10, space_after=6):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.bold = bold
    run.font.name = "Calibri"
    p.paragraph_format.space_after = Pt(space_after)
    return p


def add_code_block(doc, code):
    p = doc.add_paragraph()
    run = p.add_run(code)
    run.font.size = Pt(7.5)
    run.font.name = "Consolas"
    run.font.color.rgb = RGBColor(0, 0, 0)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)


def make_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Light Grid Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for i, h in enumerate(headers):
        set_cell(table.rows[0].cells[i], h, bold=True, size=8,
                 color=(255, 255, 255), align="center")
        shading = table.rows[0].cells[i]._element
        from docx.oxml.ns import qn
        from lxml import etree
        tc_pr = shading.find(qn("w:tcPr"))
        if tc_pr is None:
            tc_pr = etree.SubElement(shading, qn("w:tcPr"))
        shd = etree.SubElement(tc_pr, qn("w:shd"))
        shd.set(qn("w:fill"), "003366")
        shd.set(qn("w:val"), "clear")

    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            align = "right" if c_idx > 0 else "left"
            set_cell(table.rows[r_idx + 1].cells[c_idx], str(val),
                     size=8, align=align)

    return table


def main():
    doc = Document()

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10)

    for section in doc.sections:
        section.top_margin = Cm(1.5)
        section.bottom_margin = Cm(1.5)
        section.left_margin = Cm(2.0)
        section.right_margin = Cm(2.0)

    # ── TITLE ───────────────────────────────────────────────────
    title = doc.add_heading(
        "CRTPX Strategy Specification for AmiBroker Testing", level=0
    )
    for run in title.runs:
        run.font.color.rgb = RGBColor(0, 51, 102)

    add_para(doc, "Potomac Fund Management  |  Prepared by Woody Wiegmann  |  February 2026",
             size=9, space_after=2)
    add_para(doc, "Two variants of the Tactically Passive strategy for CRTPX, "
             "each with integrated tax-loss harvesting.", size=9)

    doc.add_page_break()

    # ════════════════════════════════════════════════════════════
    # STRATEGY A: PASSIVE + TLH
    # ════════════════════════════════════════════════════════════
    add_heading(doc, "Strategy A: CRTPX Passive + TLH", level=1)

    add_para(doc, "Overview", bold=True, size=11, space_after=3)
    add_para(doc,
        "Binary risk-on/risk-off toggle on the S&P 500 using a 4-indicator "
        "Penta signal. When risk-on, hold 100% VOO. When risk-off, hold 100% SGOV. "
        "Tax-loss harvesting via a 3-ticker swap ring on each side.")

    add_heading(doc, "Signal Architecture (Penta)", level=2)
    add_para(doc,
        "Four binary trend indicators, each comparing the asset's closing price "
        "to its 50-day simple moving average (SMA). A signal is GREEN if price > SMA, "
        "RED otherwise.", space_after=4)

    make_table(doc,
        ["#", "Indicator", "Ticker", "Rule"],
        [
            ["1", "S&P 500 Trend", "SPY (or ^GSPC)", "Close > 50-day SMA"],
            ["2", "Transports", "IYT (or ^DJT)", "Close > 50-day SMA"],
            ["3", "NYSE Breadth", "^NYA (or VTI)", "Close > 50-day SMA"],
            ["4", "Corporate Credit", "LQD", "Close > 50-day SMA"],
        ])

    add_para(doc, "")
    add_para(doc, "Decision Rule:", bold=True, size=10, space_after=3)
    add_para(doc, "Penta Score = sum of 4 binary indicators (0-4).\n"
             "Risk-ON if score >= 3 (3 or 4 green). Risk-OFF if score <= 2.\n"
             "3-day confirmation: the new regime must persist for 3 consecutive "
             "trading days before a switch is executed. This filters whipsaws.",
             space_after=6)

    add_heading(doc, "Allocations", level=2)
    make_table(doc,
        ["Regime", "Holding", "Weight", "Expense Ratio"],
        [
            ["Risk-On (Penta >= 3)", "VOO (Vanguard S&P 500)", "100%", "0.03%"],
            ["Risk-Off (Penta <= 2)", "SGOV (0-3mo T-Bills)", "100%", "0.03%"],
        ])

    add_heading(doc, "Tax-Loss Harvesting", level=2)
    add_para(doc,
        "When any holding has an unrealized loss exceeding 3%, sell and immediately "
        "buy the next ticker in the swap ring. Track 31-day wash sale window per ticker.",
        space_after=4)

    make_table(doc,
        ["Side", "Primary", "Swap 1", "Swap 2"],
        [
            ["Equity", "VOO", "IVV", "SPLG"],
            ["Cash", "SGOV", "BIL", "SHV"],
        ])

    add_heading(doc, "AmiBroker AFL Code", level=2)
    add_code_block(doc, """// ══════════════════════════════════════════════════════════
// CRTPX STRATEGY A: PASSIVE + TLH
// ══════════════════════════════════════════════════════════

SMA_Period = 50;
ConfirmDays = 3;

// Penta Indicators (each: price > 50-day SMA = 1, else 0)
Penta1 = Foreign("SPY","C") > MA(Foreign("SPY","C"), SMA_Period);
Penta2 = Foreign("IYT","C") > MA(Foreign("IYT","C"), SMA_Period);
Penta3 = Foreign("~NYA","C") > MA(Foreign("~NYA","C"), SMA_Period);
Penta4 = Foreign("LQD","C") > MA(Foreign("LQD","C"), SMA_Period);

PentaScore = Penta1 + Penta2 + Penta3 + Penta4;

// Raw signal
RawOn = PentaScore >= 3;

// 3-day confirmation filter
ConfirmedOn  = Sum(RawOn, ConfirmDays) == ConfirmDays;
ConfirmedOff = Sum(!RawOn, ConfirmDays) == ConfirmDays;

// Trade signals (for SPY/VOO as the traded instrument)
Buy  = ExRem(ConfirmedOn, ConfirmedOff);
Sell = ExRem(ConfirmedOff, ConfirmedOn);

// Position sizing: 100% equity when in, 0% when out
// Risk-off return: apply Foreign("SGOV","C") or set to cash
SetPositionSize(100, spsPercentOfEquity);""")

    add_heading(doc, "QuantConnect Backtest Results (June 2019 - Feb 2026)", level=2)
    make_table(doc,
        ["Metric", "Value"],
        [
            ["CAGR", "11.27%"],
            ["Total Return", "103.9%"],
            ["Max Drawdown", "-13.9%"],
            ["Sharpe Ratio", "0.54"],
            ["Sortino Ratio", "0.48"],
            ["Alpha", "0.02"],
            ["Beta (vs SPY)", "0.31"],
            ["Annual Std Dev", "9.3%"],
            ["Win Rate", "60%"],
            ["Avg Win / Avg Loss", "4.07% / -1.50%"],
            ["Regime Switches", "~32 (4.8/year)"],
            ["Total Fees ($1M)", "$5,478"],
        ])

    add_heading(doc, "Sensitivity Parameters", level=2)
    make_table(doc,
        ["Parameter", "Default", "Test Range"],
        [
            ["SMA Period", "50", "20, 50, 100, 200"],
            ["Confirmation Days", "3", "1, 2, 3, 5"],
            ["Penta Threshold", "3 of 4", "2/4, 3/4, 4/4"],
            ["TLH Loss Trigger", "-3%", "-2%, -3%, -5%"],
        ])

    doc.add_page_break()

    # ════════════════════════════════════════════════════════════
    # STRATEGY B: ENHANCED (CONVEX RISK-OFF) + TLH
    # ════════════════════════════════════════════════════════════
    add_heading(doc, "Strategy B: CRTPX Enhanced (Convex Risk-Off) + TLH", level=1)

    add_para(doc, "Overview", bold=True, size=11, space_after=3)
    add_para(doc,
        "Identical Penta signal and risk-on allocation as Strategy A. "
        "The difference: instead of parking 100% in T-bills during risk-off, "
        "we allocate to a defensive blend designed to generate positive returns "
        "during drawdowns while preserving capital.")

    add_heading(doc, "Academic Basis", level=2)
    add_para(doc,
        'Baltussen, Martens & van der Linden, "The Best Defensive Strategies: '
        'Two Centuries of Evidence," Financial Analysts Journal, Vol. 82(1), '
        "Jan 2026. Key findings over a 222-year sample (1800-2021):", space_after=4)

    add_para(doc,
        "1. Trend-following is the strongest standalone defensive strategy "
        "(4.8% return at 5% vol). Positive in BOTH up and down markets.\n\n"
        "2. DAR4020 (defensive absolute return) provides immediate put-like "
        "protection at the onset of drawdowns, before trend-following can "
        "reposition. Negative beta to 60/40.\n\n"
        "3. 50/50 blend of trend-following + DAR4020 is optimal: reduces "
        ">20% drawdowns from -37.6% to -15.1%, IR of 1.37.\n\n"
        "4. Gold is unreliable as a defensive strategy over 222 years "
        "(negative returns in worst months).\n\n"
        "5. Treasuries are historically unreliable as an equity hedge.",
        space_after=6)

    add_heading(doc, "Risk-Off Allocation (Paper-Informed)", level=2)
    add_para(doc,
        "We map the paper's findings to investable ETFs:", space_after=4)

    make_table(doc,
        ["Paper Concept", "ETF", "Weight", "Role", "Expense Ratio"],
        [
            ["Trend-following", "DBMF", "40%", "Managed futures, positive returns in crises", "0.85%"],
            ["Convexity / DAR proxy", "CAOS", "30%", "Put-like payoff, immediate drawdown protection", "0.59%"],
            ["Cash anchor", "SGOV", "30%", "Capital preservation, zero equity beta", "0.03%"],
        ])

    add_para(doc, "")
    add_para(doc, "Why not gold? The Baltussen paper shows gold has NEGATIVE "
             "returns in the worst 10% of 60/40 months over 222 years. "
             "It delivered in 1973-74 and 2007-09 but failed in most other drawdowns.",
             space_after=4)
    add_para(doc, "Why CAOS over puts? CAOS provides the convex payoff profile "
             "of a put overlay without the constant premium bleed. It maps to "
             "the paper's DAR4020 concept: immediate protection at drawdown onset.",
             space_after=6)

    add_heading(doc, "Full Allocation Table", level=2)
    make_table(doc,
        ["Regime", "Asset", "Weight"],
        [
            ["Risk-On (Penta >= 3)", "VOO", "100%"],
            ["Risk-Off (Penta <= 2)", "DBMF", "40%"],
            ["", "CAOS", "30%"],
            ["", "SGOV", "30%"],
        ])

    add_heading(doc, "AmiBroker AFL Code", level=2)
    add_code_block(doc, """// ══════════════════════════════════════════════════════════
// CRTPX STRATEGY B: ENHANCED (CONVEX RISK-OFF) + TLH
// ══════════════════════════════════════════════════════════
// Signal is IDENTICAL to Strategy A.
// Only the risk-off allocation changes.

SMA_Period = 50;
ConfirmDays = 3;

Penta1 = Foreign("SPY","C") > MA(Foreign("SPY","C"), SMA_Period);
Penta2 = Foreign("IYT","C") > MA(Foreign("IYT","C"), SMA_Period);
Penta3 = Foreign("~NYA","C") > MA(Foreign("~NYA","C"), SMA_Period);
Penta4 = Foreign("LQD","C") > MA(Foreign("LQD","C"), SMA_Period);

PentaScore = Penta1 + Penta2 + Penta3 + Penta4;
RawOn = PentaScore >= 3;
ConfirmedOn  = Sum(RawOn, ConfirmDays) == ConfirmDays;
ConfirmedOff = Sum(!RawOn, ConfirmDays) == ConfirmDays;

Buy  = ExRem(ConfirmedOn, ConfirmedOff);
Sell = ExRem(ConfirmedOff, ConfirmedOn);

SetPositionSize(100, spsPercentOfEquity);

// ── RISK-OFF PORTFOLIO (run as separate backtest or composite) ──
// When Sell fires, allocate to:
//   40% DBMF (managed futures / trend-following)
//   30% CAOS (tail-risk convexity)
//   30% SGOV (T-bills)
//
// In AmiBroker, model this as a rotational system or use
// SetForeign() to compute the blended risk-off return:
//
// RiskOffReturn = 0.40 * ROC(Foreign("DBMF","C"),1)/100
//              + 0.30 * ROC(Foreign("CAOS","C"),1)/100
//              + 0.30 * ROC(Foreign("SGOV","C"),1)/100;""")

    add_heading(doc, "QuantConnect Backtest Results (June 2021 - Feb 2026)", level=2)
    add_para(doc, "Note: Shorter backtest period due to CAOS inception (Feb 2021).",
             size=9, space_after=4)

    make_table(doc,
        ["Metric", "Strategy B (Enhanced)", "Strategy A (Passive)*"],
        [
            ["CAGR", "11.78%", "11.27%"],
            ["Total Return", "68.3%", "103.9%"],
            ["Max Drawdown", "-16.3%", "-13.9%"],
            ["Sharpe Ratio", "0.50", "0.54"],
            ["Sortino Ratio", "0.56", "0.48"],
            ["Alpha", "0.026", "0.02"],
            ["Beta (vs SPY)", "0.31", "0.31"],
            ["Win Rate", "67%", "60%"],
            ["Total Fees ($1M)", "$4,419", "$5,478"],
        ])

    add_para(doc, "*Strategy A covers June 2019-Feb 2026. Strategy B covers "
             "June 2021-Feb 2026 due to CAOS data availability. "
             "Direct comparison requires same-period test.",
             size=8, space_after=6)

    add_heading(doc, "Sensitivity Parameters (Additional for Strategy B)", level=2)
    make_table(doc,
        ["Parameter", "Default", "Test Range"],
        [
            ["DBMF Weight", "40%", "30%, 40%, 50%"],
            ["CAOS Weight", "30%", "20%, 30%, 40%"],
            ["SGOV Weight", "30%", "20%, 30%, 40%"],
            ["Alt: Replace CAOS with HEQT", "No", "Test 30% HEQT variant"],
            ["Alt: Replace CAOS with TAIL", "No", "Test 30% TAIL variant"],
        ])

    doc.add_page_break()

    # ════════════════════════════════════════════════════════════
    # TLH OPERATIONS (BOTH STRATEGIES)
    # ════════════════════════════════════════════════════════════
    add_heading(doc, "Tax-Loss Harvesting: Operational Playbook", level=1)

    add_para(doc,
        "The TLH layer is identical for both strategies. It does not change "
        "the backtest return profile -- it is purely a tax alpha overlay. "
        "It executes at the fund level on every regime transition and daily "
        "when any holding is down >3%.", space_after=6)

    add_heading(doc, "Swap Rings", level=2)
    make_table(doc,
        ["Side", "Primary", "Swap 1", "Swap 2", "Correlation"],
        [
            ["S&P 500 Equity", "VOO", "IVV", "SPLG", ">0.99"],
            ["T-Bills / Cash", "SGOV", "BIL", "SHV", ">0.99"],
        ])

    add_heading(doc, "Execution Sequence", level=2)
    add_para(doc, "On Risk-On to Risk-Off Transition:", bold=True, size=10, space_after=3)
    add_para(doc,
        "1. Screen equity holding (VOO/IVV/SPLG) for unrealized loss.\n"
        "2. If loss exists: SELL the loser first (harvest the loss).\n"
        "3. Immediately BUY defensive allocation (Strategy A: SGOV; Strategy B: DBMF/CAOS/SGOV).\n"
        "4. Record the sold ticker and wash sale expiry date (T+31 calendar days).\n"
        "5. If the signal reverses within 31 days, buy the NEXT ticker in the ring.\n",
        space_after=4)

    add_para(doc, "On Risk-Off to Risk-On Transition:", bold=True, size=10, space_after=3)
    add_para(doc,
        "1. Screen defensive holdings for unrealized loss.\n"
        "2. If loss exists: SELL the loser first (harvest).\n"
        "3. BUY equity -- use the next available ticker in the ring.\n"
        "4. Record wash sale dates.\n",
        space_after=4)

    add_para(doc, "Daily TLH Scan:", bold=True, size=10, space_after=3)
    add_para(doc,
        "1. Check current holding's unrealized P&L.\n"
        "2. If loss > 3%: SELL, BUY next in ring, record wash sale.\n"
        "3. This harvests losses mid-regime without changing exposure.",
        space_after=6)

    add_heading(doc, "Wash Sale Rules", level=2)
    add_para(doc,
        "- 31 calendar day window per ticker after a loss sale.\n"
        "- During the window, do NOT repurchase the same ticker in ANY account.\n"
        "- The swap ring has 3 tickers so there is always an available substitute.\n"
        "- After 31 days, consolidate back to primary ticker if desired.",
        space_after=6)

    doc.add_page_break()

    # ════════════════════════════════════════════════════════════
    # VALIDATION CHECKLIST
    # ════════════════════════════════════════════════════════════
    add_heading(doc, "Validation Checklist for AmiBroker Testing", level=1)

    add_para(doc,
        "1. Signal Switch Count: ~4-5 regime switches per year. If AmiBroker "
        "produces significantly more or fewer, check SMA period and confirmation lag.\n\n"
        "2. Max Drawdown: Strategy A should land near -14%, Strategy B near -16%.\n\n"
        "3. Beta: Both strategies should show beta ~0.31 vs SPY.\n\n"
        "4. Risk-off Performance: In 2022 (the worst year in-sample), Strategy B "
        "should outperform Strategy A due to DBMF's crisis alpha (+24% in 2022).\n\n"
        "5. Risk-on Performance: Both strategies should be identical in risk-on "
        "periods (both hold 100% VOO).\n\n"
        "6. Parameter Sensitivity: If results change dramatically with small "
        "parameter changes, the strategy is overfit. The Penta signal should be "
        "robust to SMA period (50 vs 100 should produce similar switch counts).",
        space_after=6)

    add_heading(doc, "Why This Isn't Overfit", level=2)
    add_para(doc,
        "- The Penta signal uses 4 well-known market breadth indicators, each "
        "with a single parameter (50-day SMA). No curve-fitting.\n"
        "- The 3-day confirmation is a standard anti-whipsaw filter, not optimized "
        "to this specific backtest period.\n"
        "- The risk-off blend in Strategy B is derived from 222 years of evidence "
        "(Baltussen et al. 2026, FAJ), not from fitting to CRTPX's short history.\n"
        "- The TLH layer has zero impact on investment returns -- it's purely operational.\n"
        "- Both strategies have only 2 free parameters: SMA period and confirmation days.",
        space_after=6)

    # Save
    out_path = os.path.join(SCRIPT_DIR, "CRTPX_Strategy_Spec_for_AmiBroker.docx")
    doc.save(out_path)
    print(f"Document saved: {out_path}")


if __name__ == "__main__":
    main()
