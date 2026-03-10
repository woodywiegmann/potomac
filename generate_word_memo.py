"""
Generates a 1-2 page Word document memo: Trade Concepts for Testing (Feb 2026)
IDEA 1: ARKK swap (CRTOX), IDEA 2: Risk-Off Convexity (Across Funds),
IDEA 3: CRTOX Universe Analysis.
"""

import os
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
OUT_DIR = r"c:\Users\WoodyWiegmann\OneDrive - PFM\Desktop\Portfolio Ideas"

NAVY = RGBColor(0x1A, 0x3C, 0x5E)
DARK = RGBColor(0x33, 0x33, 0x33)
GRAY = RGBColor(0x66, 0x66, 0x66)
RED = RGBColor(0xC0, 0x39, 0x2B)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)


def set_cell_shading(cell, hex_color):
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{hex_color}"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def set_cell_border(cell, top="single", bottom="single", left="single", right="single",
                    color="CCCCCC", size="4"):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    borders = parse_xml(
        f'<w:tcBorders {nsdecls("w")}>'
        f'  <w:top w:val="{top}" w:sz="{size}" w:space="0" w:color="{color}"/>'
        f'  <w:bottom w:val="{bottom}" w:sz="{size}" w:space="0" w:color="{color}"/>'
        f'  <w:left w:val="{left}" w:sz="{size}" w:space="0" w:color="{color}"/>'
        f'  <w:right w:val="{right}" w:sz="{size}" w:space="0" w:color="{color}"/>'
        f'</w:tcBorders>'
    )
    tcPr.append(borders)


def add_styled_table(doc, headers, rows, col_widths=None, highlight_col=None, highlight_color="E8F4E8",
                     highlight_col2=None, highlight_color2="E8F0F8", red_col=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False

    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        run.bold = True
        run.font.size = Pt(8)
        run.font.color.rgb = WHITE
        set_cell_shading(cell, "1A3C5E")
        set_cell_border(cell, color="1A3C5E")
        if col_widths:
            cell.width = Inches(col_widths[i])

    for r_idx, row_data in enumerate(rows):
        for c_idx, val in enumerate(row_data):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = ""
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if c_idx > 0 else WD_ALIGN_PARAGRAPH.LEFT
            run = p.add_run(val)
            run.font.size = Pt(8)
            if c_idx == 0:
                run.bold = True
            if red_col is not None and c_idx == red_col and r_idx > -1:
                run.font.color.rgb = RED
                run.bold = True
            if highlight_col is not None and c_idx == highlight_col:
                set_cell_shading(cell, highlight_color)
            if highlight_col2 is not None and c_idx == highlight_col2:
                set_cell_shading(cell, highlight_color2)
            set_cell_border(cell, color="CCCCCC")
            if col_widths:
                cell.width = Inches(col_widths[c_idx])

    for row in table.rows:
        row.height = Cm(0.5)
        for cell in row.cells:
            for p in cell.paragraphs:
                p.paragraph_format.space_before = Pt(1)
                p.paragraph_format.space_after = Pt(1)

    return table


def main():
    doc = Document()

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10)
    style.paragraph_format.space_after = Pt(4)
    style.paragraph_format.space_before = Pt(0)

    sections = doc.sections
    for section in sections:
        section.top_margin = Cm(1.5)
        section.bottom_margin = Cm(1.2)
        section.left_margin = Cm(2.0)
        section.right_margin = Cm(2.0)

    # ── HEADER ──
    h = doc.add_paragraph()
    h.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = h.add_run("POTOMAC FUND MANAGEMENT")
    run.bold = True
    run.font.size = Pt(16)
    run.font.color.rgb = NAVY
    h.paragraph_format.space_after = Pt(0)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = sub.add_run("Trade Concepts for Testing")
    run.bold = True
    run.font.size = Pt(13)
    run.font.color.rgb = DARK
    sub.paragraph_format.space_after = Pt(0)

    dateline = doc.add_paragraph()
    dateline.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = dateline.add_run("February 2026")
    run.font.size = Pt(11)
    run.font.color.rgb = GRAY
    dateline.paragraph_format.space_after = Pt(0)

    author = doc.add_paragraph()
    author.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = author.add_run("Woody Wiegmann")
    run.font.size = Pt(11)
    run.font.color.rgb = GRAY
    author.paragraph_format.space_after = Pt(2)

    line = doc.add_paragraph()
    line.paragraph_format.space_before = Pt(0)
    line.paragraph_format.space_after = Pt(2)
    pPr = line._p.get_or_add_pPr()
    pBdr = parse_xml(f'<w:pBdr {nsdecls("w")}><w:bottom w:val="single" w:sz="6" w:space="1" w:color="1A3C5E"/></w:pBdr>')
    pPr.append(pBdr)

    disc = doc.add_paragraph()
    run = disc.add_run("Obviously I don't have full transparency into the trading signals, so excuse any ideas that are stupid.")
    run.italic = True
    run.font.size = Pt(8.5)
    run.font.color.rgb = GRAY
    disc.paragraph_format.space_after = Pt(6)

    # ══════════════════════════════════════════════════════════
    # IDEA 1 (CRTOX)
    # ══════════════════════════════════════════════════════════
    idea1_h = doc.add_paragraph()
    run1a = idea1_h.add_run("IDEA 1 (CRTOX): ")
    run1a.bold = True
    run1a.font.size = Pt(11)
    run1a.font.color.rgb = NAVY
    run1b = idea1_h.add_run("Replace ARKK with QQQJ in CRTOX")
    run1b.bold = True
    run1b.font.size = Pt(11)
    run1b.font.color.rgb = DARK
    idea1_h.paragraph_format.space_before = Pt(4)
    idea1_h.paragraph_format.space_after = Pt(4)

    p1 = doc.add_paragraph()
    p1.paragraph_format.space_after = Pt(3)
    run = p1.add_run(
        "ARKK's discretionary management adds uncontrolled variance to CRTOX's signal-driven "
        "framework. QQQJ (Nasdaq Next Gen 100, passive, 0.15% ER) removes manager drift."
    )
    run.font.size = Pt(9)

    p1b = doc.add_paragraph()
    p1b.paragraph_format.space_after = Pt(3)
    run = p1b.add_run(
        "2023\u20132025 on CRTOX's actual trade windows: ARKK avg \u20131.93%/trade vs QQQ +1.75%/trade "
        "= 3.68% per-trade swing from manager selection alone. QQQ won 8 of 13 trades."
    )
    run.font.size = Pt(9)

    p1c = doc.add_paragraph()
    p1c.paragraph_format.space_after = Pt(3)
    run = p1c.add_run(
        "Cost savings: 60bp ER reduction (0.75% to 0.15%). Same signals, same entry/exit dates. "
        "No indicator changes required. Just a cleaner, cheaper, more predictable instrument."
    )
    run.font.size = Pt(9)

    # ══════════════════════════════════════════════════════════
    # IDEA 2 (Across Funds)
    # ══════════════════════════════════════════════════════════
    idea2_h = doc.add_paragraph()
    run2a = idea2_h.add_run("IDEA 2 (Across Funds): ")
    run2a.bold = True
    run2a.font.size = Pt(11)
    run2a.font.color.rgb = NAVY
    run2b = idea2_h.add_run("Risk-Off Convexity Enhancement")
    run2b.bold = True
    run2b.font.size = Pt(11)
    run2b.font.color.rgb = DARK
    idea2_h.paragraph_format.space_before = Pt(8)
    idea2_h.paragraph_format.space_after = Pt(4)

    p2 = doc.add_paragraph()
    p2.paragraph_format.space_after = Pt(3)
    run = p2.add_run(
        "Replace 100% SGOV during risk-off with instruments that profit from the conditions "
        "triggering our defensive posture. CAOS (tail-risk puts) + DBMF (trend-following) + "
        "SGOV (cash anchor), or HEQT (hedged equity) + DBMF + SGOV. No signal changes \u2014 "
        "only what the fund holds while waiting."
    )
    run.font.size = Pt(9)

    table_title = doc.add_paragraph()
    run = table_title.add_run("173 Risk-Off Days \u2014 Comparative Performance")
    run.bold = True
    run.font.size = Pt(9)
    run.font.color.rgb = NAVY
    table_title.paragraph_format.space_before = Pt(4)
    table_title.paragraph_format.space_after = Pt(2)

    headers = ["", "SGOV\n(current)", "50/50\nSGOV/CAOS", "EqWt\n3-Way", "15H/15D\n/70S"]
    rows = [
        ["Annualized", "+5.21%", "+6.74%", "+8.39%", "+7.92%"],
        ["Geometric", "+3.64%", "+4.71%", "+5.85%", "+5.56%"],
        ["Volatility", "0.24%", "2.56%", "4.48%", "2.55%"],
        ["Beta to S&P", "0.0005", "0.002", "0.07", "0.10"],
        ["Incremental", "\u2014", "+1.07%", "+2.21%", "+1.92%"],
    ]
    add_styled_table(doc, headers, rows,
                     col_widths=[1.1, 0.9, 0.9, 0.9, 0.9],
                     highlight_col=3, highlight_color="E8F4E8",
                     highlight_col2=4, highlight_color2="E8F0F8")

    fn = doc.add_paragraph()
    run = fn.add_run(
        "173 verified risk-off days (Mar 2023 \u2013 Feb 2026). SGOV imputed for CRDBX NAV rounding. "
        "15H/15D/70S = 15% HEQT + 15% DBMF + 70% SGOV."
    )
    run.font.size = Pt(7)
    run.font.color.rgb = GRAY
    fn.paragraph_format.space_before = Pt(2)
    fn.paragraph_format.space_after = Pt(6)

    # ══════════════════════════════════════════════════════════
    # IDEA 3 (CRTOX) - Universe Analysis
    # ══════════════════════════════════════════════════════════
    idea3_h = doc.add_paragraph()
    run3a = idea3_h.add_run("IDEA 3 (CRTOX): ")
    run3a.bold = True
    run3a.font.size = Pt(11)
    run3a.font.color.rgb = NAVY
    run3b = idea3_h.add_run("CRTOX Universe Analysis")
    run3b.bold = True
    run3b.font.size = Pt(11)
    run3b.font.color.rgb = DARK
    idea3_h.paragraph_format.space_before = Pt(8)
    idea3_h.paragraph_format.space_after = Pt(4)

    # Universe Changes
    uc_h = doc.add_paragraph()
    run = uc_h.add_run("Universe Changes")
    run.bold = True
    run.font.size = Pt(9.5)
    run.font.color.rgb = NAVY
    uc_h.paragraph_format.space_after = Pt(2)

    p3a = doc.add_paragraph()
    run = p3a.add_run("Add: ")
    run.bold = True
    run.font.size = Pt(9)
    run = p3a.add_run("URNM (uranium), COPX (copper), PAVE (infrastructure), CIBR (cybersecurity), "
                      "AMLP (midstream MLPs), EMXC (EM ex-China), XBI (equal-weight biotech)")
    run.font.size = Pt(9)
    p3a.paragraph_format.space_after = Pt(2)

    p3b = doc.add_paragraph()
    run = p3b.add_run("Drop/Limit: ")
    run.bold = True
    run.font.size = Pt(9)
    run = p3b.add_run("IGV, IAI, SOXX (high overlap with existing holdings)")
    run.font.size = Pt(9)
    p3b.paragraph_format.space_after = Pt(2)

    p3c = doc.add_paragraph()
    run = p3c.add_run("Result: ")
    run.bold = True
    run.font.size = Pt(9)
    run = p3c.add_run("Average pairwise correlation drops from 0.574 to 0.501")
    run.font.size = Pt(9)
    p3c.paragraph_format.space_after = Pt(4)

    # Backtest table
    bt_h = doc.add_paragraph()
    run = bt_h.add_run("Backtest Comparison (2021\u20132026)")
    run.bold = True
    run.font.size = Pt(9.5)
    run.font.color.rgb = NAVY
    bt_h.paragraph_format.space_after = Pt(2)

    bt_headers = ["Metric", "Current", "Proposed", "Proposed +\nNew Risk-Off"]
    bt_rows = [
        ["Ann Return", "19.59%", "20.97%", "21.28%"],
        ["Ann Vol", "20.26%", "22.26%", "22.28%"],
        ["Sharpe", "0.97", "0.94", "0.96"],
        ["Max DD", "-26.22%", "-22.95%", "-22.60%"],
        ["Calmar", "0.75", "0.91", "0.94"],
    ]
    add_styled_table(doc, bt_headers, bt_rows,
                     col_widths=[1.0, 1.0, 1.0, 1.2],
                     highlight_col=3, highlight_color="E8F4E8")

    bt_fn = doc.add_paragraph()
    bt_fn.paragraph_format.space_before = Pt(2)
    bt_fn.paragraph_format.space_after = Pt(4)

    # Recommendations
    rec_h = doc.add_paragraph()
    run = rec_h.add_run("Recommendations")
    run.bold = True
    run.font.size = Pt(9.5)
    run.font.color.rgb = NAVY
    rec_h.paragraph_format.space_after = Pt(2)

    recs = [
        "Implement expanded universe to reduce correlation and add distinct macro themes.",
        "Replace T-bill-only risk-off with SGOV/DBMF/CAOS blend for positive expected return during risk-off periods.",
        "Consider CTA (Simplify) as permanent 5\u201310% strategic allocation for systematic commodity trend exposure.",
        "Activate tax-loss harvesting protocol with correlated-but-not-identical pair ETFs.",
    ]
    for i, rec in enumerate(recs):
        p = doc.add_paragraph()
        run = p.add_run(f"{i+1}. {rec}")
        run.font.size = Pt(9)
        p.paragraph_format.space_after = Pt(1)
        p.paragraph_format.left_indent = Inches(0.15)

    # ── FOOTER ──
    footer_line = doc.add_paragraph()
    footer_line.paragraph_format.space_before = Pt(10)
    footer_line.paragraph_format.space_after = Pt(2)
    pPr = footer_line._p.get_or_add_pPr()
    pBdr = parse_xml(f'<w:pBdr {nsdecls("w")}><w:bottom w:val="single" w:sz="6" w:space="1" w:color="1A3C5E"/></w:pBdr>')
    pPr.append(pBdr)

    bl = doc.add_paragraph()
    run = bl.add_run("Bottom Line: ")
    run.bold = True
    run.font.size = Pt(9)
    run.font.color.rgb = NAVY
    run = bl.add_run(
        "All three ideas are instrument swaps and process improvements, not signal changes. "
        "Same architecture, cleaner execution, lower costs, and risk-off capital that works instead of sitting idle."
    )
    run.font.size = Pt(9)
    run.font.color.rgb = DARK

    # Save
    out_path = os.path.join(OUT_DIR, "Woody Trade Concepts - Feb 2026.docx")
    doc.save(out_path)
    print(f"Saved to: {out_path}")

    out_path2 = os.path.join(SCRIPT_DIR, "Woody Trade Concepts - Feb 2026.docx")
    doc.save(out_path2)
    print(f"Also saved to: {out_path2}")


if __name__ == "__main__":
    main()
